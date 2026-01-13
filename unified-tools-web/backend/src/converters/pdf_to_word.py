#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF文档转Word转换器
"""

import os
import uuid
from datetime import datetime
import pdfplumber
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def convert_pdf_to_word(input_file, output_file=None, options=None):
    """
    PDF转Word的主函数入口

    Args:
        input_file (str): 输入的PDF文件路径
        output_file (str, optional): 输出的Word文件路径
        options (dict, optional): 转换选项

    Returns:
        dict: 转换结果信息
    """
    if options is None:
        options = {}

    # 解析输出路径
    if output_file:
        output_dir = os.path.dirname(output_file)
        base_name = os.path.splitext(os.path.basename(output_file))[0]
    else:
        output_dir = os.path.dirname(input_file)
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}.docx")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    # 检查文件扩展名
    file_extension = os.path.splitext(input_file)[1].lower()
    
    if file_extension != '.pdf':
        raise Exception(f"不支持的文件格式: {file_extension}，仅支持PDF格式")

    # 使用pdfplumber处理PDF格式文件
    try:
        print(f"开始转换PDF文件: {input_file}")
        
        # 创建Word文档对象
        doc = Document()
        
        # 打开PDF文件
        with pdfplumber.open(input_file) as pdf:
            page_count = len(pdf.pages)
            print(f"PDF文件共有 {page_count} 页")
            
            # 遍历所有页面
            for page_num, page in enumerate(pdf.pages, 1):
                print(f"正在处理第 {page_num}/{page_count} 页")
                
                # 提取页面文本
                text = page.extract_text()
                
                if text:
                    # 添加页码标题
                    doc.add_heading(f'第 {page_num} 页', level=2)
                    
                    # 添加页面文本
                    # 按段落分割文本
                    paragraphs = text.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                # 提取页面表格
                tables = page.extract_tables()
                if tables:
                    print(f"第 {page_num} 页包含 {len(tables)} 个表格")
                    for table in tables:
                        # 创建Word表格
                        table_rows = len(table)
                        table_cols = len(table[0])
                        
                        if table_rows > 0 and table_cols > 0:
                            # 添加表格标题
                            doc.add_paragraph('表格', style='Heading 3')
                            
                            # 创建表格
                            doc_table = doc.add_table(rows=table_rows, cols=table_cols)
                            
                            # 填充表格内容
                            for i, row in enumerate(table):
                                for j, cell_value in enumerate(row):
                                    if cell_value:
                                        doc_table.cell(i, j).text = cell_value.strip()
                            
                            # 添加表格后的空行
                            doc.add_paragraph()
                
                # 提取页面图片
                # 注意：pdfplumber提取图片的功能有限，这里暂不实现图片提取
                
                # 除了最后一页，其他页面添加分页符
                if page_num < page_count:
                    doc.add_page_break()
        
        # 保存Word文档
        doc.save(output_file)
        print(f"PDF文件转换成功: {output_file}")
        
        return {
            "output_file": output_file,
            "page_count": page_count,
            "success": True
        }
    except Exception as e:
        print(f"PDF文件转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"转换PDF文件失败: {str(e)}")
