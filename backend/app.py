import logging
import os
import re
import tempfile
import threading
import time
import uuid
from typing import Optional

from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Body
from fastapi.responses import FileResponse, HTMLResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from docx import Document
import urllib.parse

from src.converters import (
    convert_docx_to_md,
    convert_markdown_to_html,
    convert_web_to_docx,
    convert_word_to_pdf,
    convert_pdf_to_word,
    convert_markdown_to_docx,
)
from src.crawlers.media_crawler import MediaCrawler

# ─── Logging ────────────────────────────────────────────────
log_level = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, log_level, logging.INFO),
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("smartdatapro")

# ─── FastAPI App ─────────────────────────────────────────────
app = FastAPI(
    title="智能文档处理平台",
    description="提供文档格式转换服务，支持多种格式转换",
    version="2.3.0",
    docs_url=os.getenv("DOCS_URL", "/docs"),  # 生产环境可设为空关闭
    redoc_url=None,
    openapi_url=os.getenv("OPENAPI_URL", "/openapi.json"),
)

# GZip 压缩
app.add_middleware(GZipMiddleware, minimum_size=1000, compresslevel=5)

# CORS（从环境变量读取允许的域名）
cors_origins = os.getenv("CORS_ORIGINS", "http://localhost:5180,http://localhost").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS", "DELETE", "PUT"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition", "Content-Length"],
)


# ─── Middleware ──────────────────────────────────────────────
@app.middleware("http")
async def add_process_time_header(request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = time.time() - start_time
    response.headers["X-Process-Time"] = str(process_time)
    response.headers["Cache-Control"] = "public, max-age=3600"
    return response


# ─── Config ──────────────────────────────────────────────────
TEMP_DIR = tempfile.gettempdir()
frontend_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend")
dist_dir = os.path.join(frontend_dir, "dist")

if os.path.exists(dist_dir):
    app.mount("/static", StaticFiles(directory=dist_dir), name="frontend")


def generate_unique_filename(original_filename: str, suffix: str) -> str:
    """生成唯一的临时文件名"""
    unique_id = uuid.uuid4().hex
    if original_filename:
        name_part = re.sub(r"[^a-zA-Z0-9_-]", "", os.path.splitext(original_filename)[0])
        return f"{name_part}_{unique_id}.{suffix}"
    return f"temp_{unique_id}.{suffix}"


def _make_error_doc(title: str, url: str, message: str) -> bytes:
    """生成转换失败的 Word 文档"""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"URL: {url}")
    doc.add_paragraph(f"错误信息: {message}")
    tmp = os.path.join(TEMP_DIR, f"error_{int(time.time())}.docx")
    doc.save(tmp)
    with open(tmp, "rb") as f:
        content = f.read()
    os.remove(tmp)
    return content


def _error_response(filename: str, content: bytes) -> Response:
    """返回错误文档响应"""
    encoded = urllib.parse.quote(filename)
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
            "Content-Length": str(len(content)),
        },
    )


# ─── Routes: Root ───────────────────────────────────────────
@app.get("/")
def root():
    index_path = os.path.join(frontend_dir, "dist/index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path, headers={"X-Content-Type-Options": "nosniff", "X-Frame-Options": "DENY"})
    return {"message": "智能文档处理平台 API", "version": "2.3.0"}


@app.get("/api/")
def read_api_root():
    return {"message": "智能文档处理平台", "version": "2.3.0"}


# ─── Routes: Conversion ─────────────────────────────────────
@app.post("/api/convert/docx-to-md")
async def convert_docx_to_md_endpoint(file: UploadFile = File(...)):
    """将Word文件转换为Markdown"""
    temp_file_path = None
    output_file = None
    try:
        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext not in [".doc", ".docx"]:
            raise HTTPException(status_code=400, detail="只支持DOC和DOCX格式文件")

        temp_file_path = os.path.join(TEMP_DIR, file.filename)
        with open(temp_file_path, "wb") as f:
            f.write(await file.read())

        result = convert_docx_to_md(temp_file_path)
        output_file = result["output_file"]

        if not os.path.exists(output_file):
            raise HTTPException(status_code=500, detail="转换失败: 生成的文件不存在")

        return FileResponse(path=output_file, filename=os.path.basename(output_file), media_type="text/markdown")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"docx-to-md conversion failed: {e}")
        raise HTTPException(status_code=500, detail=f"转换失败: {e}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)


@app.post("/api/convert/markdown-to-html")
async def convert_markdown_to_html_endpoint(file: UploadFile = File(...), style: str = Form("default")):
    """将Markdown文件转换为HTML"""
    temp_file_path = None
    output_file = None
    try:
        file_name = file.filename or "temp.md"
        if not file_name.endswith((".md", ".markdown", ".txt")):
            file_name += ".md"

        content = await file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="文件过大，最大支持10MB")

        temp_file_path = os.path.join(TEMP_DIR, file_name)
        with open(temp_file_path, "wb") as f:
            f.write(content)

        result = None
        exc = None

        def _convert():
            nonlocal result, exc
            try:
                result = convert_markdown_to_html(temp_file_path, options={"style": style})
            except Exception as e:
                exc = e

        t = threading.Thread(target=_convert, daemon=True)
        t.start()
        t.join(10)

        if t.is_alive():
            raise HTTPException(status_code=504, detail="转换超时，内容可能过于复杂")
        if exc:
            raise exc
        if not result:
            raise HTTPException(status_code=500, detail="转换失败，没有返回结果")

        output_file = result["output_file"]
        with open(output_file, "r", encoding="utf-8") as f:
            html_content = f.read()

        return HTMLResponse(content=html_content, media_type="text/html")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"markdown-to-html conversion failed: {e}")
        raise HTTPException(status_code=500, detail=f"转换失败: {e}")
    finally:
        for f in [temp_file_path, output_file]:
            if f and os.path.exists(f):
                try:
                    os.remove(f)
                except Exception:
                    pass


@app.post("/api/convert/markdown-to-docx")
async def convert_markdown_to_docx_endpoint(file: UploadFile = File(...), style: str = Form("default")):
    """将Markdown文件转换为Word"""
    temp_file_path = None
    output_file = None
    try:
        file_name = file.filename or "temp.md"
        if not file_name.endswith((".md", ".markdown", ".txt")):
            file_name += ".md"

        temp_file_path = os.path.join(TEMP_DIR, file_name)
        with open(temp_file_path, "wb") as f:
            f.write(await file.read())

        result = convert_markdown_to_docx(temp_file_path, options={"style": style, "output_dir": TEMP_DIR})
        if not result.get("success"):
            raise Exception(result.get("message", "转换失败"))

        output_file = result["output_file"]
        if not os.path.exists(output_file):
            raise HTTPException(status_code=500, detail="转换失败: 生成的文件不存在")

        return FileResponse(
            path=output_file,
            filename=os.path.basename(output_file),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"markdown-to-docx conversion failed: {e}")
        raise HTTPException(status_code=500, detail=f"转换失败: {e}")
    finally:
        for f in [temp_file_path, output_file]:
            if f and os.path.exists(f):
                try:
                    os.remove(f)
                except Exception:
                    pass


@app.post("/api/convert/web-to-docx")
async def convert_web_to_docx_endpoint(url: Optional[str] = Form(None), file: Optional[UploadFile] = File(None)):
    """将网页转换为DOCX文件"""
    temp_file_path = None
    output_file = None
    logger.info(f"web-to-docx request: url={url}, file={file.filename if file else None}")

    try:
        if not url and not file:
            raise HTTPException(status_code=400, detail="请提供URL或选择HTML文件")

        if url and not re.match(r"^https?://", url):
            raise HTTPException(status_code=400, detail="URL格式不正确，请以http://或https://开头")

        if file:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in [".html", ".htm"]:
                raise HTTPException(status_code=400, detail="只支持HTML格式文件")
            temp_file_path = os.path.join(TEMP_DIR, file.filename)
            with open(temp_file_path, "wb") as f:
                f.write(await file.read())
            url = f"file://{temp_file_path}"

        result = None
        exc = None

        def _convert():
            nonlocal result, exc
            try:
                result = convert_web_to_docx(url, options={"timeout": 10, "output_dir": TEMP_DIR})
            except Exception as e:
                exc = e

        t = threading.Thread(target=_convert, daemon=True)
        t.start()
        t.join(20)

        if t.is_alive() or exc or not result or not result.get("success"):
            msg = str(exc) if exc else result.get("message", "转换超时") if result else "转换超时"
            logger.warning(f"web-to-docx failed: {msg}")
            err_content = _make_error_doc("转换失败", url or "", msg)
            return _error_response("转换失败.docx", err_content)

        output_file = result["output_file"]
        web_title = re.sub(r'[\\/:*?"<>|]', "_", result.get("title", "网页内容"))

        return FileResponse(
            path=output_file,
            filename=f"{web_title}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"web-to-docx failed: {e}")
        err_content = _make_error_doc("转换失败", url or "", str(e))
        return _error_response("转换失败.docx", err_content)
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except Exception:
                pass


@app.post("/api/convert/pdf-to-word")
async def convert_pdf_to_word_endpoint(
    file: UploadFile = File(...), use_ocr: bool = Form(False), ocr_lang: str = Form("chi_sim+eng")
):
    """将PDF文件转换为Word文档"""
    temp_file_path = None
    output_file = None
    try:
        if not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="只支持PDF格式文件")

        temp_file_path = os.path.join(TEMP_DIR, file.filename)
        with open(temp_file_path, "wb") as f:
            f.write(await file.read())

        result = convert_pdf_to_word(temp_file_path, options={"use_ocr": use_ocr, "ocr_lang": ocr_lang})
        output_file = result["output_file"]

        if not os.path.exists(output_file):
            raise HTTPException(status_code=500, detail="转换失败: 生成的文件不存在")

        return FileResponse(path=output_file, filename=os.path.basename(output_file), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"pdf-to-word failed: {e}")
        raise HTTPException(status_code=500, detail=f"转换失败: {e}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)


@app.post("/api/convert/word-to-pdf")
async def convert_word_to_pdf_endpoint(
    file: UploadFile = File(...), use_ocr: bool = Form(False), ocr_lang: str = Form("chi_sim+eng")
):
    """将Word文件转换为PDF文档"""
    temp_file_path = None
    output_file = None
    try:
        ext = os.path.splitext(file.filename or "")[1].lower()
        if ext not in [".doc", ".docx"]:
            raise HTTPException(status_code=400, detail="只支持DOC和DOCX格式文件")

        temp_file_path = os.path.join(TEMP_DIR, file.filename)
        with open(temp_file_path, "wb") as f:
            f.write(await file.read())

        result = convert_word_to_pdf(temp_file_path, options={"use_ocr": use_ocr, "ocr_lang": ocr_lang})
        output_file = result["output_file"]

        if not os.path.exists(output_file):
            raise HTTPException(status_code=500, detail="转换失败: 生成的文件不存在")

        return FileResponse(path=output_file, filename=os.path.basename(output_file), media_type="application/pdf")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"word-to-pdf failed: {e}")
        raise HTTPException(status_code=500, detail=f"转换失败: {e}")
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)


# ─── Routes: Styles ─────────────────────────────────────────
@app.get("/api/styles")
def get_styles():
    """获取支持的HTML样式列表"""
    from src.converters.markdown_to_html import STYLES

    return {
        "styles": [{"name": k, "display_name": v["name"]} for k, v in STYLES.items()]
    }


# ─── Routes: Media Crawler ──────────────────────────────────
@app.post("/api/crawl/media")
async def crawl_media(
    platform: str = Form(...),
    url: str = Form(None),
    keyword: str = Form(None),
    post_id: str = Form(None),
):
    """自媒体平台内容采集"""
    if not any([url, keyword, post_id]):
        raise HTTPException(status_code=400, detail="至少需要提供url、keyword或post_id中的一个参数")

    crawler = MediaCrawler()
    await crawler.init_browser()
    try:
        result = await crawler.crawl(platform, url, keyword, post_id)
        return result
    except Exception as e:
        logger.error(f"Media crawl failed: {e}")
        raise HTTPException(status_code=500, detail=f"采集失败: {e}")
    finally:
        await crawler.close_browser()


# ─── Routes: Bookmarks ──────────────────────────────────────
BOOKMARKS_FILE = os.path.join(os.path.dirname(__file__), "bookmarks.json")


def _read_bookmarks():
    if not os.path.exists(BOOKMARKS_FILE):
        with open(BOOKMARKS_FILE, "w", encoding="utf-8") as f:
            f.write("[]")
    import json

    with open(BOOKMARKS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def _write_bookmarks(bookmarks):
    import json

    with open(BOOKMARKS_FILE, "w", encoding="utf-8") as f:
        json.dump(bookmarks, f, ensure_ascii=False, indent=2)


@app.get("/api/v1/bookmarks")
def get_bookmarks():
    return {"success": True, "data": _read_bookmarks()}


@app.post("/api/v1/bookmarks")
def add_bookmark(bookmark: dict = Body(...)):
    if not bookmark.get("title") or not bookmark.get("url"):
        raise HTTPException(status_code=400, detail="书签标题和URL不能为空")
    bookmarks = _read_bookmarks()
    new = {
        "id": str(uuid.uuid4()),
        "title": bookmark["title"],
        "url": bookmark["url"],
        "description": bookmark.get("description", ""),
        "createdAt": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    bookmarks.append(new)
    _write_bookmarks(bookmarks)
    return {"success": True, "data": new}


@app.put("/api/v1/bookmarks/{bookmark_id}")
def update_bookmark(bookmark_id: str, bookmark: dict = Body(...)):
    bookmarks = _read_bookmarks()
    for i, b in enumerate(bookmarks):
        if b["id"] == bookmark_id:
            bookmarks[i] = {**b, "title": bookmark.get("title", b["title"]), "url": bookmark.get("url", b["url"]), "description": bookmark.get("description", b.get("description", ""))}
            _write_bookmarks(bookmarks)
            return {"success": True, "data": bookmarks[i]}
    raise HTTPException(status_code=404, detail="书签不存在")


@app.delete("/api/v1/bookmarks/{bookmark_id}")
def delete_bookmark(bookmark_id: str):
    bookmarks = _read_bookmarks()
    new = [b for b in bookmarks if b["id"] != bookmark_id]
    if len(new) == len(bookmarks):
        raise HTTPException(status_code=404, detail="书签不存在")
    _write_bookmarks(new)
    return {"success": True, "message": "书签删除成功"}


# ─── Routes: Frontend ───────────────────────────────────────
FRONTEND_ROUTES = ["/external-md", "/markdown-editor", "/media-crawler", "/word-to-md", "/web-to-docx", "/pdf-to-word", "/word-to-pdf", "/bookmarks"]


@app.get("/{path:path}")
def frontend_catch_all(path: str):
    """所有非API请求返回index.html，由React Router处理"""
    if path.startswith("api/"):
        raise HTTPException(status_code=404, detail="API endpoint not found")
    index_path = os.path.join(frontend_dir, "dist/index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path, headers={"X-Content-Type-Options": "nosniff", "X-Frame-Options": "DENY"})
    return {"message": "智能文档处理平台 API", "version": "2.3.0"}
