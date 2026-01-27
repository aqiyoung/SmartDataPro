#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF文档转Word转换器 - 支持OCR
"""

import os
import uuid
import re
import pdfplumber
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pytesseract
from PIL import Image

import numpy as np
from io import BytesIO
from ..utils.ocr_engine import perform_ocr, check_environment


def convert_pdf_to_word(input_file, output_file=None, options=None):
    """
    PDF转Word的主函数入口

    Args:
        input_file (str): 输入的PDF文件路径
        output_file (str, optional): 输出的Word文件路径
        options (dict, optional): 转换选项
            - use_ocr: 是否使用OCR识别文字（默认True）
            - ocr_lang: OCR识别语言（默认'chi_sim+eng'）

    Returns:
        dict: 转换结果信息
    """
    if options is None:
        options = {}
    
    # 解析OCR选项，默认启用OCR
    use_ocr = options.get('use_ocr', True)
    ocr_lang = options.get('ocr_lang', 'chi_sim+eng')

    # 解析输出路径
    if output_file:
        output_dir = os.path.dirname(output_file)
        base_name = os.path.splitext(os.path.basename(output_file))[0]
    else:
        output_dir = os.path.dirname(input_file)
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}.docx")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    # 检查文件格式
    if not input_file.lower().endswith('.pdf'):
        raise Exception(f"不支持的文件格式，仅支持PDF格式")

    try:
        print(f"开始转换PDF文件: {input_file}")
        print(f"OCR设置: use_ocr={use_ocr}, ocr_lang={ocr_lang}")
        
        # 创建Word文档对象
        doc = Document()
        
        # 打开PDF文件
        with pdfplumber.open(input_file) as pdf:
            page_count = len(pdf.pages)
            print(f"PDF文件共有 {page_count} 页")
            
            # 遍历所有页面
            for page_num in range(page_count):
                print(f"正在处理第 {page_num + 1}/{page_count} 页")
                page = pdf.pages[page_num]
                
                # 1. 提取并处理文本
                text = page.extract_text()
                
                # 将整个页面转换为图像，用于OCR和图片插入
                print(f"将第 {page_num + 1} 页转换为图像")
                page_image = page.to_image(resolution=300)
                original_image = page_image.original
                
                # 保存页面图像到临时文件
                temp_img_path = f"temp_page_{page_num + 1}.png"
                original_image.save(temp_img_path, format='PNG')
                print(f"保存页面图像到: {temp_img_path}")
                
                # 2. 执行OCR获取文本
                if use_ocr:
                    print(f"对第 {page_num + 1} 页执行OCR")
                    # 使用新的 OCR 引擎，默认模式
                    ocr_text = perform_ocr(original_image, lang=ocr_lang)
                    print(f"OCR识别结果长度: {len(ocr_text)} 字符")
                    
                    if ocr_text.strip():
                        # 合并原始文本和OCR文本
                        if text:
                            combined_text = text + "\n" + ocr_text
                        else:
                            combined_text = ocr_text
                        
                        # 处理合并后的文本
                        lines = combined_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                                
                            # 跳过包含页码或图像标记的行
                            if any(keyword in line for keyword in ['第1页', '第 1 页', '[图像', '图像 1']):
                                print(f"跳过行: {line}")
                                continue
                                
                            # 清理行内的标记
                            clean_line = re.sub(r'第\s*\d+\s*页', '', line)
                            clean_line = re.sub(r'\[图像\s+\d+\]:?', '', clean_line)
                            clean_line = re.sub(r'\s+', ' ', clean_line).strip()
                            
                            if clean_line:
                                doc.add_paragraph(clean_line)
                                print(f"添加行: {clean_line}")
                else:
                    # 仅使用原始文本
                    if text and text.strip():
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                                
                            # 跳过包含页码或图像标记的行
                            if any(keyword in line for keyword in ['第1页', '第 1 页', '[图像', '图像 1']):
                                continue
                                
                            # 清理行内的标记
                            clean_line = re.sub(r'第\s*\d+\s*页', '', line)
                            clean_line = re.sub(r'\[图像\s+\d+\]:?', '', clean_line)
                            clean_line = re.sub(r'\s+', ' ', clean_line).strip()
                            
                            if clean_line:
                                doc.add_paragraph(clean_line)
                
                # 3. 插入页面图像到Word文档
                print(f"将第 {page_num + 1} 页图像插入到Word文档")
                try:
                    # 直接从临时文件插入图片
                    doc.add_picture(temp_img_path, width=Inches(5.0))
                    doc.add_paragraph()
                    print(f"成功插入第 {page_num + 1} 页图像")
                except Exception as e:
                    print(f"插入图片失败: {str(e)}")
                    import traceback
                    traceback.print_exc()
                
                # 删除临时图像文件
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                    print(f"删除临时图像文件: {temp_img_path}")
                
                # 4. 添加分页符（除了最后一页）
                if page_num < page_count - 1:
                    doc.add_page_break()
        
        # 保存Word文档
        doc.save(output_file)
        print(f"PDF文件转换成功: {output_file}")
        
        return {
            "output_file": output_file,
            "page_count": page_count,
            "success": True,
            "ocr_used": use_ocr
        }
    except Exception as e:
        print(f"PDF文件转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"转换PDF文件失败: {str(e)}")
