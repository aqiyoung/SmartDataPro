#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档转Markdown转换器
"""

import os
import uuid
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from datetime import datetime
import mammoth

SUPPORTED_IMAGE_FORMATS = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tiff",
    "image/x-emf": ".emf",
    "image/x-wmf": ".wmf",
}


class Docx2MdConverter:
    """Word文档转Markdown转换器类"""

    def __init__(
        self,
        docx_file,
        output_dir=None,
        image_dir="images",
        use_md_table=True,
        image_prefix="img_",
        keep_image_format=True,
        output_file=None,
    ):
        self.docx_file = docx_file
        self.output_dir = output_dir or os.path.dirname(docx_file)
        self.image_dir = os.path.join(self.output_dir, image_dir)
        self.use_md_table = use_md_table
        self.image_prefix = image_prefix
        self.keep_image_format = keep_image_format
        
        # 保存输出文件路径
        self.output_file = output_file
        
        # 如果没有提供输出文件路径，使用默认路径
        if not self.output_file:
            base_name = os.path.splitext(os.path.basename(self.docx_file))[0]
            self.output_file = os.path.join(self.output_dir, f"{base_name}.md")

        self.image_counter = 0
        self.table_counter = 0
        self.list_counter = 0

        self.images = {}
        self.image_paragraphs = []

        self.doc = Document(docx_file)
        self._prepare_output_dir()

    def _prepare_output_dir(self):
        """准备输出目录"""
        if not os.path.exists(self.image_dir):
            os.makedirs(self.image_dir)

    def _get_image_extension(self, content_type):
        """获取图片扩展名"""
        for ct, ext in SUPPORTED_IMAGE_FORMATS.items():
            if ct in content_type.lower():
                return ext
        return ".png"

    def _extract_images(self):
        """提取文档中的图片"""
        for rel in self.doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                try:
                    self.image_counter += 1
                    content_type = rel.target_part.content_type
                    image_data = rel.target_part.blob

                    if self.keep_image_format:
                        ext = self._get_image_extension(content_type)
                    else:
                        ext = ".png"

                    image_uuid = str(uuid.uuid4())[:8]
                    image_name = (
                        f"{self.image_prefix}{self.image_counter}_{image_uuid}{ext}"
                    )
                    image_path = os.path.join(self.image_dir, image_name)

                    with open(image_path, "wb") as f:
                        f.write(image_data)

                    self.images[rel.rId] = {
                        "name": image_name,
                        "path": image_path,
                        "relative_path": os.path.join(
                            os.path.basename(self.image_dir), image_name
                        ),
                        "content_type": content_type,
                        "size": len(image_data),
                    }
                except Exception as e:
                    print(f"图片提取错误: {e}")

    def _collect_image_paragraphs(self):
        """收集包含图片的段落"""
        for paragraph in self.doc.paragraphs:
            has_image = False
            for run in paragraph.runs:
                for drawing in run._element.findall(".//" + qn("w:drawing")):
                    has_image = True
                    break
                if has_image:
                    break
            if has_image:
                self.image_paragraphs.append(paragraph)

    def _convert_paragraph(self, paragraph):
        """转换段落为Markdown格式"""
        md_text = ""

        # 处理段落样式（标题识别）
        if paragraph.style.name.startswith("Heading"):
            try:
                level = int(paragraph.style.name.replace("Heading ", ""))
                md_text += "#" * level + " "
            except ValueError:
                pass

        # 处理段落内容
        for run in paragraph.runs:
            text = run.text.strip()
            if not text:
                continue

            # 处理文本格式
            formatted_text = text

            # 加粗
            if run.bold:
                formatted_text = f"**{formatted_text}**"

            # 斜体
            if run.italic:
                formatted_text = f"*{formatted_text}*"

            # 下划线
            if run.underline:
                formatted_text = f"<u>{formatted_text}</u>"

            # 删除线
            if run.font.strike:
                formatted_text = f"~~{formatted_text}~~"

            md_text += formatted_text

        # 处理对齐方式
        if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            md_text = f"<center>{md_text}</center>"
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            md_text = f"<div style='text-align: right;'>{md_text}</div>"

        return md_text.strip()

    def _convert_table(self, table):
        """转换表格为Markdown格式"""
        self.table_counter += 1

        if self.use_md_table:
            # 使用Markdown表格格式
            md_table = []

            # 处理表头
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            md_table.append(f"| {' | '.join(headers)} |")
            md_table.append(f"| {' | '.join(['---' for _ in headers])} |")

            # 处理表格内容
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                md_table.append(f"| {' | '.join(cells)} |")

            return "\n".join(md_table)
        else:
            # 使用HTML表格格式
            md_table = ["<table>"]

            # 处理表头
            header_row = table.rows[0]
            md_table.append("  <thead>")
            md_table.append("    <tr>")
            for cell in header_row.cells:
                md_table.append(f"      <th>{cell.text.strip()}</th>")
            md_table.append("    </tr>")
            md_table.append("  </thead>")

            # 处理表格内容
            md_table.append("  <tbody>")
            for row in table.rows[1:]:
                md_table.append("    <tr>")
                for cell in row.cells:
                    md_table.append(f"      <td>{cell.text.strip()}</td>")
                md_table.append("    </tr>")
            md_table.append("  </tbody>")
            md_table.append("</table>")

            return "\n".join(md_table)

    def _convert_list(self, paragraph):
        """转换列表为Markdown格式"""
        if paragraph.style.name.startswith("List Bullet"):
            return f"- {paragraph.text.strip()}"
        elif paragraph.style.name.startswith("List Number"):
            self.list_counter += 1
            return f"{self.list_counter}. {paragraph.text.strip()}"
        return paragraph.text.strip()

    def convert(self):
        """执行转换"""
        # 提取图片
        self._extract_images()

        # 收集包含图片的段落
        self._collect_image_paragraphs()

        # 开始转换内容
        md_content = []

        # 遍历文档内容
        for element in self.doc.element.body:
            if element.tag.endswith("p"):  # 段落
                # 查找对应的段落对象
                paragraph = None
                for p in self.doc.paragraphs:
                    if p._element == element:
                        paragraph = p
                        break

                if not paragraph:
                    continue

                # 检查是否是列表
                if paragraph.style.name.startswith("List"):
                    list_md = self._convert_list(paragraph)
                    md_content.append(list_md)
                    continue

                # 普通段落
                para_md = self._convert_paragraph(paragraph)
                if para_md:
                    md_content.append(para_md)
                    md_content.append("")

            elif element.tag.endswith("tbl"):  # 表格
                # 查找对应的表格对象
                table = None
                for t in self.doc.tables:
                    if t._element == element:
                        table = t
                        break

                if not table:
                    continue

                table_md = self._convert_table(table)
                md_content.append(table_md)
                md_content.append("")

        # 保存Markdown文件，使用self.output_file作为输出路径
        with open(self.output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))

        return {
            "output_file": self.output_file,
            "image_count": self.image_counter,
            "table_count": self.table_counter,
        }


def convert_docx_to_md(input_file, output_file=None, options=None):
    """
    Word转Markdown的主函数入口

    Args:
        input_file (str): 输入的Word文件路径
        output_file (str, optional): 输出的Markdown文件路径
        options (dict, optional): 转换选项

    Returns:
        dict: 转换结果信息
    """
    if options is None:
        options = {}

    # 解析输出路径
    if output_file:
        output_dir = os.path.dirname(output_file)
        base_name = os.path.splitext(os.path.basename(output_file))[0]
    else:
        output_dir = os.path.dirname(input_file)
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}.md")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    # 检查文件扩展名
    file_extension = os.path.splitext(input_file)[1].lower()
    
    if file_extension == '.doc':
        # 使用mammoth库处理.doc格式文件
        try:
            print(f"开始转换.doc文件: {input_file}")
            with open(input_file, "rb") as doc_file:
                print(f"正在读取.doc文件: {input_file}")
                result = mammoth.convert_to_markdown(doc_file)
                markdown_content = result.value
                print(f".doc文件转换成功，内容长度: {len(markdown_content)}字符")
                
            # 保存Markdown内容
            print(f"正在保存转换结果到: {output_file}")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            print(f"转换结果保存成功: {output_file}")
            
            return {
                "output_file": output_file,
                "image_count": 0,
                "table_count": 0,
                "success": True
            }
        except Exception as e:
            print(f".doc文件转换失败: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"转换.doc文件失败: {str(e)}")
    elif file_extension == '.docx':
        # 使用Docx2MdConverter处理.docx格式文件
        try:
            print(f"开始转换.docx文件: {input_file}")
            # 创建转换器实例
            converter = Docx2MdConverter(
                docx_file=input_file,
                output_dir=output_dir,
                image_dir=options.get("image_dir", "images"),
                use_md_table=options.get("use_md_table", True),
                image_prefix=options.get("image_prefix", "img_"),
                keep_image_format=options.get("keep_image_format", True),
                output_file=output_file,
            )

            # 执行转换
            result = converter.convert()
            result["output_file"] = output_file  # 确保返回指定的输出文件路径
            result["success"] = True
            print(f".docx文件转换成功: {output_file}")
            return result
        except Exception as e:
            print(f".docx文件转换失败: {str(e)}")
            import traceback
            traceback.print_exc()
            # 转换失败，直接抛出异常
            raise Exception(f"转换.docx文件失败: {str(e)}")
    else:
        raise Exception(f"不支持的文件格式: {file_extension}")
