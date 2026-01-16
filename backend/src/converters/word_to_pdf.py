#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档转PDF转换器 - 支持OCR
"""

import os
import uuid
import re
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
import mammoth
import platform
import pytesseract
from PIL import Image as PILImage
import numpy as np
from ..utils.ocr_engine import perform_ocr

# 注册中文字体
def register_chinese_fonts():
    # 尝试不同的中文字体路径，适配不同系统
    font_paths = []
    
    if platform.system() == 'Windows':
        font_paths.append('C:\\Windows\\Fonts\\simhei.ttf')  # 黑体
        font_paths.append('C:\\Windows\\Fonts\\simsun.ttc')  # 宋体
        font_paths.append('C:\\Windows\\Fonts\\simkai.ttf')  # 楷体
    elif platform.system() == 'Linux':
        font_paths.extend([
            '/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc',
            '/usr/share/fonts/wqy-microhei/wqy-microhei.ttc',
            '/usr/share/fonts/simsun.ttc'
        ])
    elif platform.system() == 'Darwin':  # macOS
        font_paths.extend([
            '/System/Library/Fonts/PingFang.ttc',
            '/Library/Fonts/Arial Unicode.ttf'
        ])
    
    # 注册可用的字体
    for font_path in font_paths:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont('Chinese', font_path))
                return True
            except Exception as e:
                print(f"注册字体失败: {font_path}, 错误: {e}")
    
    print("未找到可用的中文字体，中文显示可能会有问题")
    return False

# 注册中文字体
register_chinese_fonts()


def convert_word_to_pdf(input_file, output_file=None, options=None):
    """
    Word转PDF的主函数入口

    Args:
        input_file (str): 输入的Word文件路径
        output_file (str, optional): 输出的PDF文件路径
        options (dict, optional): 转换选项
            - use_ocr: 是否使用OCR识别图像中的文字（默认False）
            - ocr_lang: OCR识别语言（默认'chi_sim+eng'）

    Returns:
        dict: 转换结果信息
    """
    if options is None:
        options = {}
    
    # 解析OCR选项
    use_ocr = options.get('use_ocr', False)
    ocr_lang = options.get('ocr_lang', 'chi_sim+eng')

    # 解析输出路径
    if output_file:
        output_dir = os.path.dirname(output_file)
        base_name = os.path.splitext(os.path.basename(output_file))[0]
    else:
        output_dir = os.path.dirname(input_file)
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}.pdf")

    # 确保输出目录存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # 检查文件扩展名
    file_extension = os.path.splitext(input_file)[1].lower()
    
    if file_extension not in ['.doc', '.docx']:
        raise Exception(f"不支持的文件格式: {file_extension}，仅支持DOC和DOCX格式")

    try:
        print(f"开始转换Word文件: {input_file}")
        
        # 创建PDF文档
        doc = SimpleDocTemplate(
            output_file,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # 获取样式表
        styles = getSampleStyleSheet()
        
        # 创建自定义样式，使用中文字体
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName='Chinese',
            fontSize=18,
            leading=22,
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontName='Chinese',
            fontSize=14,
            leading=18,
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='Chinese',
            fontSize=12,
            leading=16,
        )
        
        # 文档内容列表
        story = []
        
        if file_extension == '.docx':
            # 处理.docx格式文件
            word_doc = Document(input_file)
            
            # 定义命名空间
            nsmap = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                'v': 'urn:schemas-microsoft-com:vml'
            }
            
            # 2. 遍历所有文档元素
            for element in word_doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    # 查找对应的段落对象
                    paragraph = None
                    for p in word_doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    
                    if paragraph:
                        # 检查段落中是否有文字
                        text = paragraph.text.strip()
                        if text:
                            # 跳过匹配"第X页"格式的段落
                            if re.match(r'^第\d+页$', text):
                                print(f"跳过页码段落: {text}")
                                continue
                            
                            # 检查是否是标题
                            if paragraph.style.name.startswith('Heading'):
                                try:
                                    level = int(paragraph.style.name.replace('Heading ', ''))
                                    if level == 1:
                                        p = Paragraph(text, title_style)
                                    elif level == 2:
                                        p = Paragraph(text, subtitle_style)
                                    else:
                                        p = Paragraph(text, normal_style)
                                except ValueError:
                                    p = Paragraph(text, normal_style)
                            else:
                                p = Paragraph(text, normal_style)
                            
                            story.append(p)
                            story.append(Spacer(1, 0.1 * inch))
                        
                        # 3. 处理段落中的图片
                        # 遍历段落的所有 run
                        for run in paragraph.runs:
                            # 查找所有图片引用 ID
                            image_ids = []
                            
                            # 1. 查找 drawing 中的 blip (Word 2007+)
                            blips = run._element.findall('.//a:blip', namespaces=nsmap)
                            for blip in blips:
                                embed_id = blip.get(f"{{{nsmap['r']}}}embed")
                                if embed_id:
                                    image_ids.append(embed_id)
                                    
                            # 2. 查找 VML 中的 imagedata (旧版 Word)
                            imagedatas = run._element.findall('.//v:imagedata', namespaces=nsmap)
                            for imagedata in imagedatas:
                                rel_id = imagedata.get(f"{{{nsmap['r']}}}id")
                                if rel_id:
                                    image_ids.append(rel_id)
                                    
                            # 处理找到的所有图片
                            for img_id in image_ids:
                                if img_id in word_doc.part.rels:
                                    try:
                                        rel = word_doc.part.rels[img_id]
                                        if rel.reltype == RT.IMAGE:
                                            image_data = rel.target_part.blob
                                            img_stream = BytesIO(image_data)
                                            
                                            try:
                                                # 尝试使用PIL打开图片，验证格式
                                                pil_img = PILImage.open(img_stream)
                                                # 如果不是RGB模式，转换一下（比如CMYK或P模式），reportlab可能不支持
                                                if pil_img.mode not in ('RGB', 'L'):
                                                    pil_img = pil_img.convert('RGB')
                                                    new_stream = BytesIO()
                                                    pil_img.save(new_stream, format='PNG')
                                                    img_stream = new_stream
                                                else:
                                                    # 重置流位置
                                                    img_stream.seek(0)
                                            except Exception as pil_e:
                                                print(f"PIL处理图片失败 (ID: {img_id}): {pil_e}")
                                                # 如果PIL都打不开，那reportlab肯定也挂，跳过
                                                continue
                                                
                                            img = Image(img_stream)
                                            
                                            # 计算合适的图片大小，保持比例
                                            max_width = 6.0 * inch  # 稍微放宽一点宽度
                                            # 获取图片原始尺寸
                                            img_width = img.drawWidth
                                            img_height = img.drawHeight
                                            
                                            if img_width > max_width:
                                                scale = max_width / img_width
                                                img.drawWidth = max_width
                                                img.drawHeight = img_height * scale
                                            
                                            # 确保图片高度不超过页面高度
                                            max_height = 9.0 * inch
                                            if img.drawHeight > max_height:
                                                scale = max_height / img.drawHeight
                                                img.drawWidth = img.drawWidth * scale
                                                img.drawHeight = max_height
                                            
                                            # 添加图片到PDF
                                            story.append(img)
                                            
                                            # 如果启用了OCR，对图片执行OCR
                                            if use_ocr:
                                                print(f"对图片执行OCR (ID: {img_id})")
                                                ocr_text = perform_ocr(image_data, lang=ocr_lang)
                                                if ocr_text.strip():
                                                    story.append(Paragraph("[图像OCR结果]:", normal_style))
                                                    story.append(Paragraph(ocr_text, normal_style))
                                            
                                            story.append(Spacer(1, 0.1 * inch))
                                    except Exception as e:
                                        print(f"处理图片失败 (ID: {img_id}): {str(e)}")
                                        
                elif element.tag.endswith('tbl'):  # 表格
                    # 查找对应的表格对象
                    table = None
                    for t in word_doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        story.append(Spacer(1, 0.2 * inch))
                        
                        # 转换表格数据
                        data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text.strip())
                            data.append(row_data)
                        
                        if data:
                            # 创建PDF表格
                            pdf_table = Table(data)
                            
                            # 添加表格样式
                            pdf_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 12),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            
                            story.append(pdf_table)
                            story.append(Spacer(1, 0.2 * inch))
        elif file_extension == '.doc':
            # 处理.doc格式文件，使用mammoth转换为HTML，再处理
            try:
                with open(input_file, "rb") as doc_file:
                    result = mammoth.convert_to_html(doc_file)
                    html_content = result.value
                    
                    # 简单处理HTML内容，提取文本
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    # 提取所有文本
                    text_content = soup.get_text()
                    
                    # 将文本按段落分割
                    paragraphs = text_content.split('\n')
                    
                    for para in paragraphs:
                        if para.strip():
                            p = Paragraph(para.strip(), normal_style)
                            story.append(p)
                            story.append(Spacer(1, 0.1 * inch))
            except Exception as e:
                print(f"处理.doc文件时出错: {str(e)}")
                raise Exception(f"转换.doc文件失败: {str(e)}")
        
        # 构建PDF文档
        doc.build(story)
        print(f"Word文件转换成功: {output_file}")
        
        return {
            "output_file": output_file,
            "success": True,
            "ocr_used": use_ocr
        }
    except Exception as e:
        print(f"Word文件转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"转换Word文件失败: {str(e)}")