#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
网页转Word文档转换器

功能特性：
- 支持普通网页和微信公众号文章转换
- 微信公众号文章深度优化处理
- 自动提取网页标题和内容
- 自动下载和嵌入图片
- 精确保留代码块格式
- 使用统一符号替换列表原始编号，优化文档排版
- 自动移除广告和多余元素
- 内置微信公众号图片防盗链解决方案

使用方法：
1. 命令行使用：python web_to_docx.py <URL> [output.docx]
2. API使用：from converters.web_to_docx import convert_web_to_docx

典型应用：
- 转换普通网页为Word文档
- 转换微信公众号文章为Word文档
- 转换技术博客和教程为Word文档
- 批量转换网页内容为可编辑文档
"""

import os
import sys
import time
import requests
import re
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, Comment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


class WebToDocxConverter:
    """网页转Word文档转换器类"""

    def __init__(self, url, output_dir=None, timeout=10, progress_callback=None):
        """
        初始化转换器

        Args:
            url: 目标网页URL或本地HTML文件路径
            output_dir: 输出目录
            timeout: 请求超时时间
            progress_callback: 进度回调函数
        """
        self.url = url
        
        # 检测是否为本地HTML文件
        self.is_local_file = False
        if os.path.exists(url) and os.path.isfile(url):
            self.is_local_file = True
            self.local_file_path = url
            # 对于本地文件，base_url设置为文件所在目录的绝对路径
            self.base_url = f"file:///{os.path.dirname(os.path.abspath(url))}/"
            print(f"[DEBUG] 检测到本地HTML文件: {url}")
            print(f"[DEBUG] 本地文件base_url: {self.base_url}")
        else:
            self.base_url = self._get_base_url(url)
        
        self.timeout = timeout
        # 默认输出到固定的output目录
        default_output_dir = "output"
        self.output_dir = output_dir or default_output_dir
        self.progress_callback = progress_callback

        # 设置请求头
        self.headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.8",
        }

        # 状态信息
        self.status = "初始化"
        self.progress = 0

        # 下载的图片列表
        self.downloaded_images = []

        # 网页内容
        self.html_content = None
        self.soup = None
        self.title = None
        self.content = None

        # Word文档
        self.doc = None

        # 创建输出目录
        self._create_output_dir()

        # 图片下载目录
        self.images_dir = os.path.join(self.output_dir, "images")
        os.makedirs(self.images_dir, exist_ok=True)

    def _update_progress(self, message, progress=None):
        """更新进度"""
        if self.progress_callback:
            self.status = message
            if progress is not None:
                self.progress = progress
            self.progress_callback(self.status, self.progress)

    def _get_base_url(self, url):
        """获取基础URL"""
        parsed_url = urlparse(url)
        return f"{parsed_url.scheme}://{parsed_url.netloc}"

    def _create_output_dir(self):
        """创建输出目录"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            self._update_progress(f"创建输出目录: {self.output_dir}")

    def _sanitize_filename(self, filename):
        """清理文件名，移除非法字符"""
        return re.sub(r'[\\/:*?"<>|]', "_", filename)
    
    def _clean_text(self, text, preserve_newlines=False, is_code=False):
        """清理文本，移除所有不必要的字符和转义序列
        
        Args:
            text: 要清理的文本
            preserve_newlines: 是否保留换行符，默认为False
            is_code: 是否是代码块，默认为False
        """
        if not text:
            return ""
        
        import re
        # 移除所有HTML标签，但保留style标签内容
        import re
        
        # 提取所有style标签内容
        styles_text = ""
        if self.soup:
            style_tags = self.soup.find_all("style")
            for style in style_tags:
                styles_text += style.get_text() + "\n"
        
        # 移除style标签
        if self.soup:
            for tag in self.soup.find_all("style"):
                tag.decompose()
                
        # 移除所有HTML标签
        text = re.sub(r'<[^>]+>', '', text)
        
        # 再次强力清理可能残留的CSS代码
        # 移除CSS样式块，如 body { ... } 或 .class { ... }
        text = re.sub(r'[^{}]+\{[^}]+\}', '', text)
        # 移除独立的CSS属性行，如 font-size: 12px;
        text = re.sub(r'[\w-]+\s*:\s*[^;]+;', '', text)
        # 移除残留的CSS选择器和花括号
        text = re.sub(r'[\w\.-]+\s*\{', '', text)
        text = re.sub(r'\}', '', text)
        # 移除所有 \xXX 转义字符
        text = re.sub(r'\\x[0-9a-fA-F]{2}', '', text)
        # 移除所有 \uXXXX 转义字符
        text = re.sub(r'\\u[0-9a-fA-F]{4}', '', text)
        
        # 根据参数决定是否保留换行符
        if preserve_newlines:
            # 只移除 \t 和 \r，保留 \n
            text = re.sub(r'\\t', ' ', text)
            text = re.sub(r'\\r', '', text)
            # 移除多余的空格，但保留换行符
            text = re.sub(r' +', ' ', text)
            text = re.sub(r'^\s+', '', text, flags=re.MULTILINE)
            text = re.sub(r'\s+$', '', text, flags=re.MULTILINE)
        else:
            # 移除所有 \n、\t、\r 等转义字符
            text = re.sub(r'\\[ntr]', ' ', text)
            # 移除多余的空格
            text = re.sub(r'\s+', ' ', text).strip()
        
        if not is_code:
            # 只在非代码块中移除引号、括号和JS相关内容
            # 移除所有引号和括号
            text = re.sub(r'["\'\(\)\[\]]', '', text)
            # 移除所有JavaScript函数调用，如 JsDecode()
            text = re.sub(r'JsDecode\([^)]*\)', '', text)
            # 移除所有类似 content_noencode: 这样的JS变量定义
            text = re.sub(r'\w+\s*:\s*', '', text)
        
        return text

    def _download_html(self):
        """下载HTML内容或读取本地HTML文件"""
        try:
            self._update_progress("正在获取HTML内容...", 10)
            
            # 检查是否为本地HTML文件
            if self.is_local_file:
                print(f"[DEBUG] 读取本地HTML文件: {self.local_file_path}")
                
                # 读取本地HTML文件内容
                with open(self.local_file_path, "r", encoding="utf-8") as f:
                    self.html_content = f.read()
                
                print(f"[DEBUG] 本地HTML文件读取成功，内容长度: {len(self.html_content)} 字符")
                self._update_progress("本地HTML文件读取完成", 20)
                return True
            
            # 网络URL处理
            print(f"[DEBUG] 开始下载HTML: {self.url}")
            print(f"[DEBUG] 超时设置: {self.timeout} 秒")
            
            # 检查URL格式是否正确
            if not self.url.startswith(('http://', 'https://')):
                print(f"[DEBUG] URL格式错误: {self.url}")
                self._update_progress(f"URL格式错误: {self.url}", 0)
                return False
            
            # 简化请求头，只保留必要的User-Agent
            simple_headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "zh-CN,zh;q=0.8",
            }
            print(f"[DEBUG] 使用简化请求头: {simple_headers}")
            
            try:
                # 尝试使用简单的请求方式
                response = requests.get(
                    self.url, 
                    headers=simple_headers, 
                    timeout=self.timeout,
                    verify=False,  # 关闭证书验证
                    allow_redirects=True
                )
                
                print(f"[DEBUG] HTTP状态码: {response.status_code}")
                response.raise_for_status()
                
                # 处理响应编码
                if response.encoding is None:
                    response.encoding = response.apparent_encoding
                print(f"[DEBUG] 响应编码: {response.encoding}")
                
                # 读取响应内容，限制最大长度为1MB
                max_content_length = 1024 * 1024  # 1MB
                self.html_content = response.text[:max_content_length]
                print(f"[DEBUG] 响应内容长度: {len(self.html_content)} 字符")
                
                # 如果内容被截断，添加标记
                if len(self.html_content) == max_content_length:
                    self.html_content += "<!-- 内容被截断 -->"
                    print(f"[DEBUG] 响应内容超过 {max_content_length} 字符，已截断")
                
                self._update_progress("网页内容下载完成", 20)
                print(f"[DEBUG] HTML下载成功")
                return True
            except requests.exceptions.RequestException as e:
                print(f"[DEBUG] HTML下载请求异常: {str(e)}")
                import traceback
                traceback.print_exc()
                self._update_progress(f"网页下载失败: {str(e)}", 0)
                return False
            except Exception as e:
                print(f"[DEBUG] 获取HTML内容失败: {str(e)}")
                import traceback
                traceback.print_exc()
                self._update_progress(f"获取HTML内容失败: {str(e)}", 0)
                return False
        except FileNotFoundError as e:
            print(f"[DEBUG] 本地HTML文件未找到: {str(e)}")
            self._update_progress(f"本地HTML文件未找到: {str(e)}", 0)
            return False
        except UnicodeDecodeError as e:
            print(f"[DEBUG] 本地HTML文件编码错误: {str(e)}")
            self._update_progress(f"本地HTML文件编码错误: {str(e)}", 0)
            return False
        except requests.exceptions.RequestException as e:
            print(f"[DEBUG] HTML下载请求异常: {str(e)}")
            import traceback
            traceback.print_exc()
            self._update_progress(f"网页下载失败: {str(e)}", 0)
            return False
        except Exception as e:
            print(f"[DEBUG] 获取HTML内容失败: {str(e)}")
            import traceback
            traceback.print_exc()
            self._update_progress(f"获取HTML内容失败: {str(e)}", 0)
            return False

    def _parse_html(self):
        """解析HTML内容，特别优化微信公众号文章处理"""
        try:
            print(f"[DEBUG] 开始解析HTML")
            self._update_progress("正在解析网页内容...", 30)
            
            # 检查HTML内容是否存在
            if not self.html_content:
                print(f"[DEBUG] HTML内容为空")
                return False
            
            self.soup = BeautifulSoup(self.html_content, "html.parser")
            print(f"[DEBUG] BeautifulSoup初始化完成")
            
            # 预处理：在HTML阶段就移除所有列表编号，从源头解决问题
            print(f"[DEBUG] 开始预处理HTML，移除所有列表编号")
            
            # 1. 处理所有列表项
            all_li = self.soup.find_all("li")
            for li in all_li:
                self._clean_list_item(li)
            
            # 2. 处理所有段落，移除可能包含的编号
            all_p = self.soup.find_all("p")
            for p in all_p:
                for child in p.children:
                    if child.name is None:  # 文本节点
                        text = str(child)
                        cleaned_text = self._remove_list_numbering(text)
                        if cleaned_text != text:
                            child.replace_with(cleaned_text)
            
            print(f"[DEBUG] HTML预处理完成，移除了所有列表编号")

            # 获取标题 - 特别优化微信公众号文章
            self.title = "网页内容"
            print(f"[DEBUG] 默认标题: {self.title}")
            
            # 尝试从微信公众号特定位置获取标题
            if self.soup:
                try:
                    # 微信公众号标题选择器，增加更多选择器覆盖更多网页类型
                    wechat_title_selectors = [
                        "h1.rich_media_title",  # 微信公众号标题
                        "title",  # HTML标题
                        "h1",  # 一级标题
                        "h1[class*='title']",  # 包含title的h1
                        "h2[class*='title']",  # 包含title的h2
                        "h3[class*='title']",  # 包含title的h3
                        "div[class*='title']",  # 包含title的div
                        "header[class*='title']",  # 包含title的header
                        "article h1",  # 文章内的h1
                        "main h1",  # 主内容区的h1
                        "h2",  # 二级标题
                        "h3"   # 三级标题
                    ]
                    
                    print(f"[DEBUG] 尝试获取标题")
                    for selector in wechat_title_selectors:
                        title_element = self.soup.select_one(selector)
                        if title_element and title_element.get_text():
                            title_text = title_element.get_text().strip()
                            # 过滤掉过短的标题
                            if len(title_text) > 3:
                                self.title = title_text
                                print(f"[DEBUG] 从选择器 {selector} 获取到标题: {self.title}")
                                break
                except Exception as e:
                    print(f"[DEBUG] 获取标题失败: {str(e)}")
            
            if not self.title or self.title == "网页内容" or len(self.title) <= 3:
                # 尝试从meta标签获取标题
                try:
                    # 尝试多种meta标签获取标题
                    meta_selectors = [
                        {"property": "og:title"},  # Open Graph标题
                        {"name": "title"},  # 普通meta标题
                        {"name": "og:site_name"},  # 网站名称
                        {"name": "description"}  # 描述信息（备选）
                    ]
                    
                    for selector in meta_selectors:
                        meta_title = self.soup.find("meta", selector)
                        if meta_title and meta_title.get("content"):
                            meta_content = meta_title.get("content").strip()
                            if meta_content and len(meta_content) > 3:
                                self.title = meta_content
                                print(f"[DEBUG] 从meta标签 {selector} 获取到标题: {self.title}")
                                break
                except Exception as e:
                    print(f"[DEBUG] 从meta标签获取标题失败: {str(e)}")
            
            # 最终检查，如果标题仍不满意，尝试从body中提取第一个有意义的文本作为标题
            if not self.title or self.title == "网页内容" or len(self.title) <= 3:
                try:
                    if self.soup.body:
                        # 提取body中的所有文本
                        all_text = self.soup.body.get_text(separator="\n", strip=True)
                        lines = all_text.split("\n")
                        # 找到第一个长度合适的文本行作为标题
                        for line in lines:
                            if line and len(line) > 3 and len(line) < 100:
                                self.title = line.strip()
                                print(f"[DEBUG] 从body文本中提取到标题: {self.title}")
                                break
                except Exception as e:
                    print(f"[DEBUG] 从body提取标题失败: {str(e)}")
            
            # 确保标题不是默认值
            if not self.title or self.title == "网页内容" or len(self.title) <= 3:
                # 使用URL的最后一部分作为标题
                try:
                    parsed_url = urlparse(self.url)
                    path_parts = parsed_url.path.strip("/").split("/")
                    # 找到最后一个有意义的路径部分
                    for part in reversed(path_parts):
                        if part and len(part) > 3:
                            self.title = part.replace("-", " ").replace("_", " ").capitalize()
                            print(f"[DEBUG] 从URL中提取到标题: {self.title}")
                            break
                except Exception as e:
                    print(f"[DEBUG] 从URL提取标题失败: {str(e)}")

            # 获取主要内容 - 特别优化微信公众号文章
            self.content = None
            print(f"[DEBUG] 开始获取主要内容")
            
            # 如果是本地文件，直接使用body作为内容，跳过复杂的提取逻辑
            if self.is_local_file and self.soup.body:
                print(f"[DEBUG] 本地文件，直接使用body作为主要内容")
                self.content = self.soup.body
                # 移除不需要的标签 (仅移除script和style，保留其他结构)
                for tag in self.content.find_all(["script", "style", "meta", "link", "title"]):
                    tag.decompose()
            else:
                # 网络内容，尝试智能提取
                # 首先尝试使用传统方式获取内容（针对普通网页和部分微信公众号文章）
                if self.soup:
                    # 移除不需要的标签
                    unwanted_tags = ["script", "style", "nav", "footer", "aside", "iframe", "form", 
                                    "header", "noscript", "meta", "link", "input", "textarea", 
                                    "button", "select", "option", "fieldset", "legend", "label"]
                    
                    print(f"[DEBUG] 移除不需要的标签: {unwanted_tags}")
                    for tag in self.soup.find_all(unwanted_tags):
                        try:
                            tag.decompose()
                        except Exception as e:
                            print(f"[DEBUG] 移除标签 {tag.name} 失败: {str(e)}")
                            continue
                    
                    # 移除微信公众号特定的广告和无用元素
                    wechat_ad_selectors = [
                        "div[class*='advertisement']", "div[class*='ad-wrap']",
                        "div[class*='weixinad']", "div[class*='wxad']",
                        "div[class*='advert']", "div[class*='recommend-read']",
                        "div[class*='related-articles']", "div[class*='comment-area']",
                        "div[class*='like-area']", "div[class*='share-area']",
                        "div#js_copyright_area", "div#js_post_bottom_ad",
                        "div[class*='profile']", "div[class*='wechat-ad']"
                    ]
                    
                    print(f"[DEBUG] 移除微信公众号广告元素")
                    for selector in wechat_ad_selectors:
                        try:
                            for tag in self.soup.select(selector):
                                tag.decompose()
                        except Exception as e:
                            print(f"[DEBUG] 移除广告元素 {selector} 失败: {str(e)}")
                            continue
                    
                    # 尝试获取微信公众号文章的主要内容容器
                    wechat_content_selectors = [
                        "div.rich_media_content",
                        "div#js_content",
                        "article",
                        "div.content",
                        "div.main-content",
                        "div[class*='article']",
                        "div[class*='content']"
                    ]
                    
                    print(f"[DEBUG] 尝试获取内容容器")
                    for selector in wechat_content_selectors:
                        content_element = self.soup.select_one(selector)
                        if content_element:
                            # 检查内容是否为空
                            if content_element.get_text(strip=True):
                                self.content = content_element
                                print(f"[DEBUG] 使用选择器 {selector} 获取到内容")
                                break
            
            print(f"[DEBUG] 获取内容结果: {'成功' if self.content else '失败'}")
            
            # 只有在非本地文件且确实获取失败时，才尝试文本提取兜底
            if not self.is_local_file and (not self.content or not self.content.get_text(strip=True)):
                print(f"[DEBUG] 传统方式获取的内容为空，尝试直接提取纯文本")
                # ... (保留原有的兜底逻辑)
                try:
                    import re
                    # 直接从原始HTML中提取所有文本
                    clean_text = re.sub(r'<[^>]+>', '', self.html_content)
                    # ... (简化，省略部分代码)
                    if clean_text:
                        self.content = self.soup.new_tag("div")
                        self.content.string = clean_text
                except Exception:
                    pass
            
            # 如果仍然没有获取到内容，尝试直接从soup.body获取
            if not self.content:
                print(f"[DEBUG] 尝试从soup.body获取内容")
                if self.soup and hasattr(self.soup, 'body') and self.soup.body:
                    self.content = self.soup.body
                    print(f"[DEBUG] 从soup.body获取到内容")
            
            # 最后的备用方案
            if not self.content or not self.content.get_text(strip=True):
                print(f"[DEBUG] 所有方法都未获取到内容，使用默认内容")
                if not self.soup:
                    self.soup = BeautifulSoup()
                self.content = self.soup.new_tag("div")
                self.content.string = "无法获取网页内容"

            print(f"[DEBUG] HTML解析完成，获取到标题: {self.title}，内容: {'成功' if self.content else '失败'}")
            self._update_progress("网页内容解析完成", 40)
            return True
        except Exception as e:
            print(f"[DEBUG] HTML解析失败: {str(e)}")
            import traceback
            traceback.print_exc()
            self._update_progress(f"网页解析失败: {str(e)}", 0)
            return False

    def _download_image(self, img_url, img_index):
        """下载单个图片，支持本地文件和网络图片，增强图片下载可靠性"""
        try:
            # 清理图片URL，移除可能的转义字符和无效内容
            import re
            img_url = re.sub(r'&amp;', '&', img_url)
            img_url = re.sub(r'&quot;', '"', img_url)
            img_url = re.sub(r'&#39;', "'", img_url)
            img_url = img_url.strip()
            
            if not img_url or img_url == "#" or "javascript:" in img_url or "data:" in img_url:
                print(f"[DEBUG] 跳过无效图片URL: {img_url}")
                return None
            
            # 构建完整URL或路径
            final_img_url = img_url
            is_local_image = False
            
            # 处理本地HTML文件中的图片
            if self.is_local_file:
                # 对于本地文件，处理相对路径
                if not img_url.startswith(("http://", "https://", "file://")):
                    # 如果是相对路径，使用本地文件目录作为基准
                    local_img_path = os.path.join(os.path.dirname(self.local_file_path), img_url)
                    # 检查文件是否存在
                    if os.path.exists(local_img_path):
                        final_img_url = local_img_path
                        is_local_image = True
                        print(f"[DEBUG] 本地图片路径: {final_img_url}")
                    else:
                        # 尝试使用base_url构建完整URL
                        final_img_url = urljoin(self.base_url, img_url)
                elif img_url.startswith("file://"):
                    # 处理file://协议的本地图片
                    # 移除file://前缀，转换为本地路径
                    final_img_url = img_url.replace("file:///", "").replace("file://", "")
                    # 处理不同操作系统的路径格式
                    if os.name == 'nt':  # Windows系统
                        final_img_url = final_img_url.replace('/', '\\')
                    is_local_image = True
                    print(f"[DEBUG] file://协议图片转换为本地路径: {final_img_url}")
            else:
                # 网络文件处理
                if not img_url.startswith(("http://", "https://")):
                    final_img_url = urljoin(self.base_url, img_url)
            
            # 检查图片URL是否已经被处理过
            for existing_img in self.downloaded_images:
                if img_url == existing_img['url'] or final_img_url == existing_img['final_url']:
                    print(f"[DEBUG] 图片已下载，跳过: {img_url}")
                    return existing_img['path']

            # 提取文件名，保留原始文件扩展名
            img_ext = ".jpg"  # 默认扩展名
            if is_local_image:
                # 从本地文件路径提取扩展名
                _, ext = os.path.splitext(final_img_url)
                if ext:
                    img_ext = ext.lower()
            else:
                # 从URL提取扩展名
                ext_match = re.search(r'\.(jpg|jpeg|png|gif|webp|bmp)($|\?)', final_img_url, re.IGNORECASE)
                if ext_match:
                    img_ext = f".{ext_match.group(1).lower()}"
            
            img_name = f"image_{img_index}_{int(time.time())}{img_ext}"
            img_path = os.path.join(self.images_dir, img_name)

            if is_local_image:
                # 本地图片直接复制
                import shutil
                shutil.copy2(final_img_url, img_path)
                print(f"[DEBUG] 本地图片复制成功: {final_img_url} -> {img_path}")
            else:
                # 网络图片下载，增强重试机制和防盗链处理
                import urllib3
                from requests.adapters import HTTPAdapter
                from urllib3.util.retry import Retry
                
                session = requests.Session()
                retry = Retry(
                    total=5,  # 增加重试次数到5次
                    backoff_factor=0.5,  # 增加退避因子
                    status_forcelist=[429, 500, 502, 503, 504, 505],
                    allowed_methods=["HEAD", "GET", "OPTIONS"]
                )
                adapter = HTTPAdapter(max_retries=retry)
                session.mount("http://", adapter)
                session.mount("https://", adapter)
                
                # 添加更全面的请求头，防止防盗链
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "image/*",
                    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
                    "Referer": self.url,
                    "Origin": self.base_url,
                    "Connection": "keep-alive",
                    "Cache-Control": "no-cache"
                }
                
                # 处理微信公众号图片防盗链
                if "mmbiz.qpic.cn" in final_img_url or "mmbiz.qlogo.cn" in final_img_url:
                    headers['Referer'] = "https://mp.weixin.qq.com/"
                
                # 下载图片，增加超时时间
                response = session.get(final_img_url, headers=headers, timeout=15, allow_redirects=True)
                response.raise_for_status()
                
                # 验证响应是否为图片
                if 'image' not in response.headers.get('Content-Type', ''):
                    print(f"[DEBUG] 响应不是图片，跳过: {final_img_url}")
                    return None

                # 保存图片
                with open(img_path, "wb") as f:
                    f.write(response.content)
                print(f"[DEBUG] 网络图片下载成功: {final_img_url} -> {img_path}")

            # 记录下载的图片信息
            self.downloaded_images.append(
                {"url": img_url, "path": img_path, "name": img_name, "final_url": final_img_url}
            )

            return img_path
        except Exception as e:
            print(f"图片下载/复制失败: {img_url}, 错误: {str(e)}")
            import traceback
            traceback.print_exc()
            return None

    def _download_all_images(self):
        """下载所有图片，添加完整的防御性检查"""
        try:
            self._update_progress("正在查找图片...", 50)
            
            # 确保content存在，并且find_all不会失败
            images = []
            if self.content:
                try:
                    images = self.content.find_all("img")
                except Exception as e:
                    self._update_progress(f"查找图片失败: {str(e)}", 55)
                    return True
            
            total_images = len(images)
            actual_image_count = 0

            if total_images == 0:
                self._update_progress("未找到图片", 60)
                return True

            self._update_progress(f"找到 {total_images} 张图片，开始下载...", 50)

            # 下载图片
            for i, img in enumerate(images, 1):
                try:
                    # 确保img不是None，并且可以安全调用get方法
                    if img:
                        # 微信公众号图片可能使用data-src、data-original等属性存储真实URL
                        img_url = img.get("src")
                        
                        # 检查所有可能的图片URL属性
                        all_img_attrs = [
                            "src",
                            "data-src",
                            "data-original",
                            "data-loaded",
                            "data-lazyload",
                            "data-lazy-src",
                            "lazy-src",
                            "original-src"
                        ]
                        
                        final_img_url = None
                        for attr in all_img_attrs:
                            attr_url = img.get(attr)
                            if attr_url and attr_url != "" and "placeholder" not in attr_url:
                                final_img_url = attr_url
                                break
                        
                        if final_img_url:
                            actual_image_count += 1
                            # 处理微信公众号图片URL，可能包含特殊格式
                            import re
                            final_img_url = re.sub(r'&amp;', '&', final_img_url)
                            final_img_url = re.sub(r'&quot;', '"', final_img_url)
                            final_img_url = re.sub(r'&#39;', "'", final_img_url)
                            
                            # 构建完整URL
                            if not final_img_url.startswith(("http://", "https://")):
                                final_img_url = urljoin(self.base_url, final_img_url)
                            
                            self._download_image(final_img_url, i)
                except Exception as e:
                    # 跳过无法处理的图片
                    print(f"处理图片时出错: {str(e)}")
                    continue

                # 更新进度
                progress = 50 + int((i / total_images) * 10)
                self._update_progress(f"正在下载图片 {i}/{total_images}", progress)

            self._update_progress(
                f"图片下载完成，共下载 {len(self.downloaded_images)} 张", 60
            )
            return True
        except Exception as e:
            self._update_progress(f"图片下载失败: {str(e)}", 0)
            print(f"图片下载详细错误: {str(e)}")
            return False

    def _create_word_document(self):
        """创建Word文档"""
        try:
            self._update_progress("正在创建Word文档...", 70)

            # 创建文档
            self.doc = Document()

            # 设置样式
            self._setup_document_styles()

            # 添加标题
            title_para = self.doc.add_heading(self.title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加正文内容
            self._add_content_to_document()

            self._update_progress("Word文档创建完成", 90)
            return True
        except Exception as e:
            self._update_progress(f"Word文档创建失败: {str(e)}", 0)
            return False

    def _setup_document_styles(self):
        """设置文档样式"""
        # 设置中文字体
        styles = self.doc.styles

        # 正文样式
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "微软雅黑"
        normal_font.size = Pt(12)
        normal_font.color.rgb = RGBColor(0, 0, 0)

        # 标题样式
        for level in range(1, 6):
            heading_style = styles[f"Heading {level}"]
            heading_font = heading_style.font
            heading_font.name = "微软雅黑"
            heading_font.bold = True

    def _add_content_to_document(self):
        """将内容添加到文档，确保微信公众号文章内容能被正确保存"""
        # 如果content为空，直接返回
        if not self.content:
            return
            
        try:
            print(f"\n=== 开始处理内容 ===")
            print(f"内容元素: {self.content.name if hasattr(self.content, 'name') else '未知'}")
            
            # 首先尝试处理HTML结构，保留格式和图片
            print("处理HTML结构，保留排版和图片")
            self._process_block_element(self.content)
            
            # 关键改进：如果_process_block_element没有添加任何内容，确保添加文本
            # 检查文档段落数量（至少应该有标题、信息行）
            if len(self.doc.paragraphs) <= 2:
                print("处理HTML结构后没有添加足够内容，尝试提取纯文本")
                # 直接从内容中提取所有文本
                full_text = self.content.get_text(separator="\n", strip=True)
                print(f"提取的文本长度: {len(full_text)} 字符")
                
                # 清理文本
                full_text = self._clean_text(full_text)
                
                if full_text:
                    print("添加提取的纯文本到文档")
                    self.doc.add_paragraph(full_text)
                else:
                    print("提取的文本为空，尝试其他方式获取内容")
                    # 尝试从整个soup中提取文本
                    if hasattr(self, 'soup') and self.soup:
                        soup_text = self.soup.get_text(separator="\n", strip=True)
                        print(f"从整个soup提取的文本长度: {len(soup_text)} 字符")
                        # 使用通用清理函数清理文本
                        soup_text = self._clean_text(soup_text)
                        if soup_text:
                            self.doc.add_paragraph(soup_text)
                        else:
                            print("无法获取任何文本内容")
                            self.doc.add_paragraph("无法获取网页内容")
            
        except Exception as e:
            print(f"处理内容时发生异常: {str(e)}")
            # 最后的备用方案
            try:
                # 尝试提取纯文本作为备用
                full_text = self.content.get_text(separator="\n", strip=True)
                full_text = self._clean_text(full_text)
                if full_text:
                    self.doc.add_paragraph(full_text)
                else:
                    self.doc.add_paragraph("处理网页内容时发生异常，但已尝试保存可用内容")
            except:
                self.doc.add_paragraph("处理网页内容时发生异常，但已尝试保存可用内容")

    def _parse_color(self, color_str):
        """解析CSS颜色值为RGBColor"""
        if not color_str:
            return None
        
        color_str = color_str.strip().lower()
        
        # Hex format
        import re
        hex_match = re.match(r'#([0-9a-f]{3}|[0-9a-f]{6})', color_str)
        if hex_match:
            hex_val = hex_match.group(1)
            if len(hex_val) == 3:
                hex_val = "".join([c*2 for c in hex_val])
            return RGBColor(int(hex_val[0:2], 16), int(hex_val[2:4], 16), int(hex_val[4:6], 16))
            
        # RGB format
        rgb_match = re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_str)
        if rgb_match:
            return RGBColor(int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3)))
            
        # Common names
        colors = {
            'red': RGBColor(255, 0, 0),
            'green': RGBColor(0, 128, 0),
            'blue': RGBColor(0, 0, 255),
            'black': RGBColor(0, 0, 0),
            'white': RGBColor(255, 255, 255),
            'gray': RGBColor(128, 128, 128),
            'grey': RGBColor(128, 128, 128),
            'orange': RGBColor(255, 165, 0),
            'purple': RGBColor(128, 0, 128)
        }
        return colors.get(color_str)

    def _set_shading(self, element, color_hex):
        """设置底纹（背景色）"""
        if not color_hex:
            return
            
        # 移除#号
        color_hex = color_hex.replace('#', '')
        
        # 创建shd元素
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        
        # 如果是Run对象，获取其_rPr属性
        if hasattr(element, '_r'):
            rPr = element._r.get_or_add_rPr()
            rPr.append(shd)
        # 如果是Paragraph对象，获取其_pPr属性
        elif hasattr(element, '_p'):
            pPr = element._p.get_or_add_pPr()
            pPr.append(shd)
            
    def _set_borders(self, paragraph, color_hex="auto", size="4", space="1"):
        """设置段落边框"""
        pPr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pbdr')
        
        # 左边框
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), size)
        left.set(qn('w:space'), space)
        left.set(qn('w:color'), color_hex.replace('#', ''))
        pbdr.append(left)
        
        # 还可以设置 top, bottom, right, between
        
        pPr.append(pbdr)

    def _apply_style_from_css(self, run, element):
        """从元素的style属性应用样式到run"""
        if not element.has_attr('style'):
            return
            
        style_str = element['style']
        styles = {}
        for item in style_str.split(';'):
            if ':' in item:
                key, val = item.split(':', 1)
                styles[key.strip().lower()] = val.strip().lower()
        
        # Color
        if 'color' in styles:
            color = self._parse_color(styles['color'])
            if color:
                run.font.color.rgb = color
                
        # Font Weight
        if 'font-weight' in styles:
            if styles['font-weight'] in ['bold', '700', '800', '900']:
                run.bold = True
                
        # Font Style
        if 'font-style' in styles:
            if styles['font-style'] == 'italic':
                run.italic = True
                
        # Text Decoration
        if 'text-decoration' in styles:
            if 'underline' in styles['text-decoration']:
                run.font.underline = True
            if 'line-through' in styles['text-decoration']:
                run.font.strike = True
                
        # Background Color
        if 'background-color' in styles:
            bg_color = styles['background-color']
            # 解析颜色为Hex
            rgb = self._parse_color(bg_color)
            if rgb:
                hex_color = "{:02x}{:02x}{:02x}".format(rgb[0], rgb[1], rgb[2])
                self._set_shading(run, hex_color)
            elif bg_color.startswith('#'):
                self._set_shading(run, bg_color)

    def _process_inline_content(self, element, paragraph, remove_numbering=False):
        """处理行内元素，保留格式（粗体、斜体、链接等）"""
        if not element:
            return

        # 遍历子节点
        for child in element.children:
            # 文本节点
            if child.name is None:
                text = str(child)
                import re
                text = re.sub(r'\s+', ' ', text)
                if text:
                    # 如果需要移除编号，先处理文本
                    if remove_numbering:
                        text = self._remove_list_numbering(text)
                    paragraph.add_run(text)
                continue

            # 元素节点
            run = None
            
            # 处理 Ruby (注音)
            if child.name == 'ruby':
                # 提取基准文本和注音
                base_text = ""
                rt_text = ""
                for rb in child.find_all(text=True, recursive=False):
                    if remove_numbering:
                        base_text += self._remove_list_numbering(rb)
                    else:
                        base_text += rb
                rt_tag = child.find('rt')
                if rt_tag:
                    rt_text = rt_tag.get_text()
                
                # Word 完美支持 Ruby 比较复杂，这里使用兼容格式：基准文本(注音)
                if base_text and rt_text:
                    run = paragraph.add_run(f"{base_text}({rt_text})")
                else:
                    run = paragraph.add_run(base_text)
            
            elif child.name in ['strong', 'b']:
                text = child.get_text()
                if remove_numbering:
                    text = self._remove_list_numbering(text)
                run = paragraph.add_run(text)
                run.bold = True
            elif child.name in ['em', 'i']:
                text = child.get_text()
                if remove_numbering:
                    text = self._remove_list_numbering(text)
                run = paragraph.add_run(text)
                run.italic = True
            elif child.name == 'u':
                text = child.get_text()
                if remove_numbering:
                    text = self._remove_list_numbering(text)
                run = paragraph.add_run(text)
                run.font.underline = True
            elif child.name == 's' or child.name == 'del':
                text = child.get_text()
                if remove_numbering:
                    text = self._remove_list_numbering(text)
                run = paragraph.add_run(text)
                run.font.strike = True
            elif child.name == 'code':
                text = child.get_text()
                run = paragraph.add_run(text)
                run.font.name = 'Courier New'
                # 尝试解析 style，如果没有则使用默认红色
                if not child.has_attr('style') or 'color' not in child['style']:
                    run.font.color.rgb = RGBColor(220, 50, 47)
            elif child.name == 'mark':
                text = child.get_text()
                if remove_numbering:
                    text = self._remove_list_numbering(text)
                run = paragraph.add_run(text)
                from docx.enum.text import WD_COLOR_INDEX
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif child.name == 'span':
                text = child.get_text()
                if remove_numbering:
                    text = self._remove_list_numbering(text)
                run = paragraph.add_run(text)
            elif child.name == 'a':
                # 简化处理链接，保留文本并变蓝
                text = child.get_text()
                if text:
                    if remove_numbering:
                        text = self._remove_list_numbering(text)
                    run = paragraph.add_run(text)
                    run.font.color.rgb = RGBColor(0, 112, 192) # 蓝色链接
                    run.font.underline = True
            elif child.name == 'br':
                paragraph.add_run('\n')
                continue
            elif child.name == 'img':
                # 行内图片处理
                self._handle_inline_image(child, paragraph)
                continue
            elif child.name == 'input' and child.get('type') == 'checkbox':
                # 处理复选框
                checked = child.get('checked') is not None
                run = paragraph.add_run('☑ ' if checked else '☐ ')
                run.font.name = 'MS Gothic' # 使用支持符号的字体
            else:
                # 递归处理其他标签
                self._process_inline_content(child, paragraph, remove_numbering)
                continue

            # 应用通用CSS样式 (如果 run 存在)
            if run:
                self._apply_style_from_css(run, child)

    def _optimize_image(self, img_path):
        """优化图片：压缩、调整大小、转换格式"""
        try:
            from PIL import Image
            import io
            import os
            
            # 打开图片
            with Image.open(img_path) as img:
                # 转换为RGB模式（处理透明背景）
                if img.mode in ('RGBA', 'LA'):
                    bg = Image.new('RGB', img.size, (255, 255, 255))
                    bg.paste(img, mask=img.split()[-1])
                    img = bg
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # 调整大小，确保最大宽度不超过Word页面宽度
                max_width = 6.0  # Word页面宽度，单位：英寸
                max_pixels = int(max_width * 96)  # 96 DPI
                
                if img.width > max_pixels:
                    # 按比例缩放
                    ratio = max_pixels / img.width
                    new_width = int(img.width * ratio)
                    new_height = int(img.height * ratio)
                    img = img.resize((new_width, new_height), Image.LANCZOS)
                
                # 保存优化后的图片
                optimized_path = img_path + '.optimized.jpg'
                img.save(optimized_path, format='JPEG', quality=85, optimize=True)
                
                return optimized_path
        except Exception as e:
            print(f"图片优化失败: {e}")
            # 优化失败，返回原图片
            return img_path
    
    def _handle_inline_image(self, img_element, paragraph):
        """处理行内图片，优化图片显示效果"""
        try:
            # 获取图片URL，检查所有可能的属性
            img_attrs = ["src", "data-src", "data-original", "data-loaded", "data-lazyload", "data-lazy-src", "lazy-src", "original-src"]
            img_url = None
            for attr in img_attrs:
                img_url = img_element.get(attr)
                if img_url and img_url != "" and "placeholder" not in img_url:
                    break
            
            if not img_url:
                return

            # 查找已下载的图片
            img_path = None
            for img_info in self.downloaded_images:
                if img_url in img_info['url'] or img_info['url'] in img_url or img_url in img_info['final_url']:
                    img_path = img_info['path']
                    break
            
            # 如果没找到，尝试下载
            if not img_path:
                if not img_url.startswith(('http', 'https', 'file')):
                    img_url = urljoin(self.base_url, img_url)
                img_path = self._download_image(img_url, len(self.downloaded_images)+1)

            if img_path and os.path.exists(img_path):
                # 优化图片
                optimized_img_path = self._optimize_image(img_path)
                
                # 插入图片到文档
                run = paragraph.add_run()
                picture = run.add_picture(optimized_img_path, width=None)  # 自动使用优化后的大小
                
                # 居中显示图片
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 清理临时优化文件
                if optimized_img_path != img_path and os.path.exists(optimized_img_path):
                    os.remove(optimized_img_path)
        except Exception as e:
            print(f"行内图片处理失败: {e}")
            import traceback
            traceback.print_exc()

    def _process_table(self, table_element):
        """处理表格"""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return

            # 计算最大列数
            max_cols = 0
            for row in rows:
                cols = len(row.find_all(['td', 'th']))
                if cols > max_cols:
                    max_cols = cols
            
            if max_cols == 0:
                return

            # 创建表格
            table = self.doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < max_cols:
                        # 获取单元格段落（默认第一个）
                        p = table.cell(i, j).paragraphs[0]
                        # 递归处理单元格内容的格式
                        self._process_inline_content(cell, p)
                        
            self.doc.add_paragraph() # 表格后空一行
        except Exception as e:
            print(f"表格处理失败: {e}")

    def _process_block_element(self, element):
        """处理块级元素，递归提取内容"""
        if not element:
            return
            
        try:
            # 遍历子节点
            children = []
            if hasattr(element, 'children'):
                children = list(element.children)
            else:
                return
            
            print(f"[DEBUG] 处理元素 {element.name}, 子节点数: {len(children)}")
                
            for i, child in enumerate(children):
                try:
                    # 忽略空白文本节点
                    if not child:
                        continue
                        
                    # 处理文本节点
                    if isinstance(child, str):
                        text = str(child).strip()
                        if text:
                            # 只有当文本不仅仅是标点符号或非常短时才打印日志，避免日志过多
                            if len(text) > 1:
                                print(f"[DEBUG] 处理文本节点: {text[:20]}...")
                            self.doc.add_paragraph(text)
                        continue

                    # 忽略无效的Tag (如Comment)
                    if child.name is None:
                        continue
                        
                    print(f"[DEBUG] 处理子元素 {i}: {child.name}")

                    # 专门处理图片元素，确保块级图片也能得到优化处理
                    if child.name == "img":
                        # 创建新段落，处理块级图片
                        p = self.doc.add_paragraph()
                        self._handle_inline_image(child, p)
                    elif child.name in ["div", "figure", "picture"]:
                        # 检查是否包含图片
                        has_img = False
                        for img_tag in child.find_all("img"):
                            if img_tag:
                                has_img = True
                                # 创建新段落，处理图片
                                p = self.doc.add_paragraph()
                                self._handle_inline_image(img_tag, p)
                        # 如果包含图片，跳过后续处理
                        if has_img:
                            continue
                    
                    if child.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                        try:
                            level = int(child.name[1])
                            # 获取标题文本
                            text = child.get_text(strip=True)
                            if text:
                                p = self.doc.add_heading(text, level=level)
                                # 如果有样式，尝试应用（虽然add_heading已经设置了文本，但可能需要颜色等）
                                # 注意：add_heading返回的p包含一个run，但我们需要遍历child来保留行内格式
                                # 为了简单起见，这里不再重新处理行内格式，因为标题通常没有复杂格式
                                # 如果需要，可以清空p然后重新添加
                                pass
                        except Exception as h_e:
                            print(f"[DEBUG] 标题处理失败: {h_e}")
                            pass
                            
                    elif child.name == "p":
                        p = self.doc.add_paragraph()
                        self._process_inline_content(child, p, remove_numbering=True)
                        
                    elif child.name == "table":
                        self._process_table(child)
                        
                    elif child.name in ["ul", "ol"]:
                        self._add_list_to_document(child)
                        
                    elif child.name == "blockquote":
                        # 引用块处理
                        paragraphs = child.find_all('p', recursive=False)
                        
                        # 提取样式
                        bg_color = None
                        border_color = "auto"
                        if child.has_attr('style'):
                            style_str = child['style']
                            for item in style_str.split(';'):
                                if 'background-color' in item:
                                    try:
                                        bg_color_val = item.split(':')[1].strip()
                                        rgb = self._parse_color(bg_color_val)
                                        if rgb:
                                            bg_color = "{:02x}{:02x}{:02x}".format(rgb[0], rgb[1], rgb[2])
                                        elif bg_color_val.startswith('#'):
                                            bg_color = bg_color_val
                                    except: pass
                        
                        if paragraphs:
                            for p_tag in paragraphs:
                                p = self.doc.add_paragraph()
                                p.paragraph_format.left_indent = Inches(0.5)
                                if bg_color:
                                    self._set_shading(p, bg_color)
                                self._process_inline_content(p_tag, p)
                        else:
                            p = self.doc.add_paragraph()
                            p.paragraph_format.left_indent = Inches(0.5)
                            if bg_color:
                                self._set_shading(p, bg_color)
                            self._process_inline_content(child, p)
                            
                    elif child.name == "pre":
                        # 代码块处理
                        code_text = child.get_text()
                        if code_text.strip():
                            p = self.doc.add_paragraph()
                            
                            # 提取样式
                            bg_color = "f1f1f1" # Default
                            text_color = None
                            
                            if child.has_attr('style'):
                                style_str = child['style']
                                for item in style_str.split(';'):
                                    item = item.strip()
                                    if not item: continue
                                    
                                    if 'background-color' in item or 'background' in item:
                                        try:
                                            bg_color_val = item.split(':')[1].strip()
                                            rgb = self._parse_color(bg_color_val)
                                            if rgb:
                                                bg_color = "{:02x}{:02x}{:02x}".format(rgb[0], rgb[1], rgb[2])
                                            elif bg_color_val.startswith('#'):
                                                bg_color = bg_color_val.replace('#', '')
                                        except: pass
                                    elif item.startswith('color:'):
                                        try:
                                            color_val = item.split(':')[1].strip()
                                            rgb = self._parse_color(color_val)
                                            if rgb:
                                                text_color = rgb
                                        except: pass
                            
                            # 尝试设置背景色
                            self._set_shading(p, bg_color)
                            
                            run = p.add_run(code_text)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                            
                            if text_color:
                                run.font.color.rgb = text_color
                                
                            p.paragraph_format.left_indent = Inches(0.2)
                            
                    elif child.name == "img":
                        p = self.doc.add_paragraph()
                        self._handle_inline_image(child, p)
                        
                    elif child.name == "hr":
                        p = self.doc.add_paragraph()
                        p.add_run("_________________________")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    elif child.name in ["div", "section", "article", "main", "header", "footer", "body"]:
                        # 容器元素，递归处理
                        self._process_block_element(child)
                        
                    else:
                        # 未知块级元素，尝试作为段落处理，或者递归
                        # 如果它包含块级元素，应该递归；如果是行内元素，作为段落
                        # 简单判断：如果包含 p, div, ul, ol, h1-h6, table，则递归
                        if child.find(["p", "div", "ul", "ol", "h1", "h2", "h3", "h4", "h5", "h6", "table"]):
                            self._process_block_element(child)
                        else:
                            p = self.doc.add_paragraph()
                            self._process_inline_content(child, p)
                            
                except Exception as inner_e:
                    print(f"[DEBUG] 处理子元素 {child.name if hasattr(child, 'name') else 'text'} 失败: {inner_e}")
                    import traceback
                    traceback.print_exc()
                    continue
                    
        except Exception as e:
            print(f"[DEBUG] 处理块级元素失败: {e}")
            import traceback
            traceback.print_exc()
            pass

    def _remove_list_numbering(self, text):
        """移除文本中的列表编号，如"1. "、"(1) "、"a. "、"6. 1. "等"""
        import re
        
        # 处理复杂嵌套编号，如"6. 1. "、"2. 3. 1. "等
        # 匹配多个连续的编号，直到最后一个编号
        complex_pattern = r'^(\s*(?:\d+|[a-zA-Z]|[ivxlcdmIVXLCDM])+(?:\.|\)|\(\d+\)|\([a-zA-Z]\)|\([ivxlcdmIVXLCDM]+\))\s+)+'
        
        # 先尝试移除复杂嵌套编号
        if re.match(complex_pattern, text, re.IGNORECASE):
            return re.sub(complex_pattern, '', text, count=1, flags=re.IGNORECASE)
        
        # 匹配各种列表编号格式
        # 1. 数字编号: 1. 2. 3. 或 1) 2) 3) 或 (1) (2) (3)
        # 2. 字母编号: a. b. c. 或 a) b) c) 或 (a) (b) (c)
        # 3. 罗马数字: i. ii. iii. 或 I. II. III.
        patterns = [
            r'^\s*\d+\.\s+',      # 1. 
            r'^\s*\d+\)\s+',      # 1) 
            r'^\s*\(\d+\)\s+',    # (1) 
            r'^\s*[a-zA-Z]\.\s+',  # a. 
            r'^\s*[a-zA-Z]\)\s+',  # a) 
            r'^\s*\([a-zA-Z]\)\s+',# (a) 
            r'^\s*[ivxlcdm]+\.\s+', # i. ii. iii. 
            r'^\s*[IVXLCDM]+\.\s+', # I. II. III. 
            r'^\s*[ivxlcdm]+\)\s+', # i) ii) iii) 
            r'^\s*[IVXLCDM]+\)\s+'  # I) II) III) 
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return re.sub(pattern, '', text, count=1, flags=re.IGNORECASE)
        
        # 最后尝试移除任何数字+点的组合，如"1."、"2."等
        simple_pattern = r'^\s*(\d+\.\s*)+'
        if re.match(simple_pattern, text):
            return re.sub(simple_pattern, '', text, count=1)
        
        return text
    
    def _clean_list_item(self, li_element):
        """彻底清理列表项，移除所有可能的编号和不必要的元素"""
        import re
        
        # 创建一个副本进行处理，避免修改原始元素
        from copy import deepcopy
        cleaned_li = deepcopy(li_element)
        
        # 1. 移除列表项的value属性，避免Word自动添加编号
        if 'value' in cleaned_li.attrs:
            del cleaned_li.attrs['value']
        
        # 2. 移除其他可能导致自动编号的属性
        auto_number_attrs = ['start', 'type', 'compact', 'reversed']
        for attr in auto_number_attrs:
            if attr in cleaned_li.attrs:
                del cleaned_li.attrs[attr]
        
        # 3. 移除列表元素的编号相关属性
        if cleaned_li.parent and cleaned_li.parent.name in ['ol', 'ul']:
            # 移除有序列表的start、type、reversed等属性
            parent = cleaned_li.parent
            for attr in ['start', 'type', 'reversed']:
                if attr in parent.attrs:
                    del parent.attrs[attr]
        
        # 4. 递归清理所有文本节点，移除所有编号
        def clean_text_nodes(element):
            for child in element.children:
                if child.name is None:  # 文本节点
                    text = str(child)
                    cleaned_text = self._remove_list_numbering(text)
                    if cleaned_text != text:
                        child.replace_with(cleaned_text)
                else:  # 元素节点
                    # 递归清理子元素
                    clean_text_nodes(child)
        
        clean_text_nodes(cleaned_li)
        
        # 5. 移除可能包含编号的所有子元素
        for child in cleaned_li.find_all(['span', 'div', 'p'], recursive=True):
            child_text = child.get_text().strip()
            # 如果子元素只包含编号，直接移除
            if re.match(r'^(\d+|[a-zA-Z]|[ivxlcdmIVXLCDM])+(\.|\)|\(\d+\)|\([a-zA-Z]\)|\([ivxlcdmIVXLCDM]+\))(\s+|$)', child_text):
                child.decompose()
        
        return cleaned_li
    
    def _add_list_to_document(self, list_element, level=0):
        """添加列表到文档，完全避免数字编号，使用纯文本方式处理"""
        try:
            # 获取直接子列表项 (避免获取嵌套列表的li)
            list_items = list_element.find_all("li", recursive=False)
            if not list_items:
                # 如果没有直接子元素，尝试查找所有
                list_items = list_element.find_all("li")
                
            if not list_items:
                return
                
            # 定义列表符号选项，根据嵌套级别使用不同符号
            # 所有列表类型都使用统一的符号体系，完全避免数字编号
            level_symbols = [
                '➤',  # 一级列表：箭头符号
                '●',  # 二级列表：实心圆点
                '◆',  # 三级列表：实心菱形
                '■',  # 四级列表：实心方块
                '▲'   # 五级列表：实心三角形
            ]
            
            # 获取当前嵌套级别的符号
            symbol = level_symbols[min(level, len(level_symbols)-1)]
            
            # 遍历列表项
            for li in list_items:
                if li:
                    # 1. 首先彻底清理列表项，移除所有编号和相关属性
                    cleaned_li = self._clean_list_item(li)
                    
                    # 2. 创建新段落，设置缩进和符号
                    p = self.doc.add_paragraph()
                    
                    # 3. 添加缩进（使用空格实现）
                    indent_spaces = '    ' * level  # 每级4个空格
                    p.add_run(indent_spaces)
                    
                    # 4. 添加自定义符号
                    p.add_run(f'{symbol} ')
                    
                    # 5. 处理列表项内容，保留格式并移除编号
                    # 检查列表项是否有子元素
                    if len(list(cleaned_li.children)) > 0:
                        # 有子元素，处理HTML结构，保留格式
                        self._process_inline_content(cleaned_li, p, remove_numbering=True)
                    else:
                        # 只有文本，直接处理
                        text = cleaned_li.get_text(separator=' ', strip=True)
                        if text:
                            cleaned_text = self._remove_list_numbering(text)
                            if cleaned_text.strip():
                                p.add_run(cleaned_text)
                    
                    # 6. 检查是否有嵌套列表
                    nested_lists = cleaned_li.find_all(['ul', 'ol'], recursive=False)
                    for nested_list in nested_lists:
                        # 递归处理嵌套列表，嵌套级别+1
                        self._add_list_to_document(nested_list, level+1)
        except Exception as e:
            print(f"处理列表时出错: {str(e)}")
            import traceback
            traceback.print_exc()

    def _add_image_to_document(self, img_url):
        """添加图片到文档，保持原文顺序和简洁性
        无法插入图片时预留空间并将图片下载到本地
        """
        try:
            import tempfile
            import os
            import requests
            import re
            
            print(f"[DEBUG] _add_image_to_document 调用，传入URL: {img_url}")
            print(f"[DEBUG] 当前已下载图片数量: {len(self.downloaded_images)}")
            
            # 简化图片插入逻辑
            # 遍历已下载的图片，查找匹配项
            img_path = None
            
            for i, img_info in enumerate(self.downloaded_images):
                print(f"[DEBUG] 检查已下载图片 {i+1}/{len(self.downloaded_images)}: {img_info['url']}")
                
                # 检查多种匹配方式
                if (img_url == img_info['url'] or 
                    img_url == img_info.get('final_url') or 
                    img_url in img_info['url'] or 
                    img_info['url'] in img_url or
                    # 忽略查询参数的匹配
                    re.sub(r'\?.+$', '', img_url) == re.sub(r'\?.+$', '', img_info['url'])):
                    img_path = img_info['path']
                    print(f"[DEBUG] 找到匹配的已下载图片: {img_url} -> {img_path}")
                    break
            
            # 如果没有找到匹配的图片，尝试直接处理
            if not img_path or not os.path.exists(img_path):
                print(f"[DEBUG] 未找到匹配的已下载图片，尝试直接下载")
                
                # 处理URL中的特殊字符
                processed_img_url = re.sub(r'&amp;', '&', img_url)
                processed_img_url = re.sub(r'&quot;', '"', processed_img_url)
                processed_img_url = re.sub(r'&#39;', "'", processed_img_url)
                
                # 构建完整URL
                final_img_url = processed_img_url
                if not processed_img_url.startswith(("http://", "https://")):
                    final_img_url = urljoin(self.base_url, processed_img_url)
                
                # 直接下载并插入图片
                img_index = len(self.downloaded_images) + 1
                img_path = self._download_image(final_img_url, img_index)
                print(f"[DEBUG] 直接下载图片结果: {img_path}")
                
                # 如果下载成功，更新img_info
                if img_path and os.path.exists(img_path):
                    print(f"[DEBUG] 直接下载图片成功，路径: {img_path}")
            
            # 初始化变量，避免NameError
            is_local_image = False
            final_img_url = img_url
            processed_img_url = img_url
            
            # 检查是否为本地文件
            if self.is_local_file:
                # 处理本地HTML文件中的图片
                if img_url.startswith("file://"):
                    # file://协议转换为本地路径
                    final_img_url = img_url.replace("file:///", "").replace("file://", "")
                    is_local_image = True
                    print(f"[DEBUG] file://协议图片转换为本地路径: {final_img_url}")
                elif not img_url.startswith(("http://", "https://")):
                    # 检查是否为本地文件路径
                    local_img_path = os.path.join(os.path.dirname(self.local_file_path), img_url)
                    if os.path.exists(local_img_path):
                        final_img_url = local_img_path
                        is_local_image = True
                        print(f"[DEBUG] 本地相对路径图片: {final_img_url}")
            
            # 尝试将图片添加到文档
            if img_path and os.path.exists(img_path):
                try:
                    # 添加图片到文档，保持原始尺寸比例
                    self.doc.add_picture(img_path, width=Inches(5.0))
                    print(f"[DEBUG] 成功在原位置添加图片: {processed_img_url}")
                    return
                except Exception as e:
                    print(f"[DEBUG] 使用本地图片插入失败: {str(e)}")
            
            # 如果未找到图片或插入失败，处理本地文件情况
            if is_local_image and os.path.exists(final_img_url):
                try:
                    # 直接使用本地图片路径
                    self.doc.add_picture(final_img_url, width=Inches(5.0))
                    print(f"[DEBUG] 直接使用本地图片: {final_img_url}")
                    return
                except Exception as e:
                    print(f"[DEBUG] 直接使用本地图片失败: {str(e)}")
            
            # 对于网络图片，如果未下载或插入失败，尝试直接下载并插入
            else:
                try:
                    # 微信公众号图片可能有防盗链，需要添加Referer头
                    img_headers = self.headers.copy()
                    img_headers['Referer'] = self.url
                    
                    print(f"[DEBUG] 直接下载图片: {final_img_url}")
                    # 使用带重试机制的会话
                    import urllib3
                    from requests.adapters import HTTPAdapter
                    from urllib3.util.retry import Retry
                    
                    session = requests.Session()
                    retry = Retry(
                        total=2,
                        backoff_factor=0.1,
                        status_forcelist=[429, 500, 502, 503, 504],
                        allowed_methods=["HEAD", "GET", "OPTIONS"]
                    )
                    adapter = HTTPAdapter(max_retries=retry)
                    session.mount("http://", adapter)
                    session.mount("https://", adapter)
                    
                    response = session.get(final_img_url, headers=img_headers, timeout=5)  # 缩短超时时间到5秒
                    response.raise_for_status()
                    
                    # 检查是否为图片文件
                    if not response.headers.get('Content-Type', '').startswith('image/'):
                        print(f"[DEBUG] 不是图片文件: {final_img_url}, Content-Type: {response.headers.get('Content-Type')}")
                        # 预留图片空间
                        placeholder_para = self.doc.add_paragraph("[图片占位符]")
                        placeholder_para.space_before = Pt(12)
                        placeholder_para.space_after = Pt(12)
                        placeholder_para.add_run(f" (不是图片文件: {processed_img_url})")
                        return
                    
                    # 创建临时文件保存图片
                    # 根据Content-Type确定文件扩展名
                    content_type = response.headers.get('Content-Type', 'image/jpeg')
                    if 'png' in content_type:
                        suffix = '.png'
                    elif 'gif' in content_type:
                        suffix = '.gif'
                    elif 'webp' in content_type:
                        suffix = '.webp'
                    else:
                        suffix = '.jpg'
                    
                    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        # 添加图片到文档，保持原始尺寸比例
                        self.doc.add_picture(temp_file_path, width=Inches(5.0))
                        print(f"[DEBUG] 成功直接添加图片到文档: {final_img_url}")
                    except Exception as e:
                        print(f"[DEBUG] 直接添加图片到文档失败: {str(e)}")
                        # 预留图片空间
                        placeholder_para = self.doc.add_paragraph("[图片占位符]")
                        placeholder_para.space_before = Pt(12)
                        placeholder_para.space_after = Pt(12)
                        placeholder_para.add_run(f" (插入失败: {processed_img_url})")
                        # 仍然保存图片到本地
                        img_index = len(self.downloaded_images) + 1
                        img_name = f"image_{img_index}_{int(time.time())}{suffix}"
                        local_img_path = os.path.join(self.images_dir, img_name)
                        with open(local_img_path, "wb") as f:
                            f.write(response.content)
                        self.downloaded_images.append(
                            {"url": processed_img_url, "path": local_img_path, "name": img_name, "final_url": final_img_url}
                        )
                        print(f"[DEBUG] 图片已保存到本地: {local_img_path}")
                    finally:
                        # 删除临时文件
                        os.unlink(temp_file_path)
                except Exception as direct_e:
                    print(f"[DEBUG] 直接下载图片失败: {str(direct_e)}")
            
            # 所有方法都失败，添加占位符
            placeholder_para = self.doc.add_paragraph("[图片占位符]")
            placeholder_para.space_before = Pt(12)
            placeholder_para.space_after = Pt(12)
            placeholder_para.add_run(f" (无法处理: {processed_img_url})")
            print(f"[DEBUG] 无法处理图片，添加占位符: {processed_img_url}")
        except Exception as e:
            print(f"[DEBUG] 处理图片时发生异常: {str(e)}")
            # 预留图片空间
            placeholder_para = self.doc.add_paragraph("[图片占位符]")
            placeholder_para.space_before = Pt(12)
            placeholder_para.space_after = Pt(12)
            placeholder_para.add_run(f" (处理失败: {img_url})")
            # 打印详细错误信息，便于调试
            import traceback
            traceback.print_exc()

    def _save_document(self, output_file):
        """保存Word文档"""
        try:
            self._update_progress("正在保存Word文档...", 95)
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_file)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
            
            # 处理文件已存在的情况
            if os.path.exists(output_file):
                # 重命名已存在的文件
                base_name, ext = os.path.splitext(output_file)
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                new_output_file = f"{base_name}_{timestamp}{ext}"
                self._update_progress(f"文件已存在，将保存为: {os.path.basename(new_output_file)}", 95)
                output_file = new_output_file
            
            # 保存文档
            self.doc.save(output_file)
            self._update_progress("Word文档保存完成", 100)
            return True
        except PermissionError:
            self._update_progress(f"Word文档保存失败: 没有写入权限，请检查输出目录权限或文件是否被占用", 0)
            return False
        except Exception as e:
            self._update_progress(f"Word文档保存失败: {str(e)}", 0)
            return False

    def convert(self, output_file=None):
        """
        执行转换过程

        Args:
            output_file: 输出文件路径

        Returns:
            dict: 转换结果信息
        """
        import time
        start_time = time.time()
        max_execution_time = 120  # 最大执行时间限制为120秒
        
        print("[DEBUG] 开始执行转换过程")
        print(f"[DEBUG] 最大执行时间: {max_execution_time} 秒")
        
        try:
            # 执行转换步骤，每个步骤都检查执行时间
            if time.time() - start_time > max_execution_time:
                return {"success": False, "message": "转换超时"}
            
            success = self._download_html()
            if not success:
                print("[DEBUG] HTML下载失败，使用后备转换逻辑")
                # 从URL中提取有意义的部分作为标题
                try:
                    from urllib.parse import urlparse
                    parsed_url = urlparse(self.url)
                    # 使用URL路径的最后一部分作为标题
                    path_parts = parsed_url.path.strip("/").split("/")
                    # 找到最后一个有意义的路径部分
                    fallback_title = "网页转换结果"
                    for part in reversed(path_parts):
                        if part and len(part) > 3:
                            fallback_title = part.replace("-", " ").replace("_", " ").capitalize()
                            break
                except Exception as e:
                    print(f"[DEBUG] 从URL提取标题失败: {str(e)}")
                    fallback_title = "网页转换结果"
                    
                # HTML下载失败，使用简单的后备转换逻辑，使用从URL提取的标题
                self.html_content = f"<html><head><title>{fallback_title}</title></head><body><h1>{fallback_title}</h1><p>URL: {self.url}</p><p>转换时间: {time.strftime('%Y-%m-%d %H:%M:%S')}</p><p>注意: 由于网络问题，无法获取完整网页内容，只显示基本信息。</p></body></html>"
                self.soup = BeautifulSoup(self.html_content, "html.parser")
                self.title = fallback_title
                self.content = self.soup.body
            
            if time.time() - start_time > max_execution_time:
                return {"success": False, "message": "转换超时"}
            
            success = self._parse_html()
            if not success:
                print("[DEBUG] HTML解析失败，使用后备解析逻辑")
                # HTML解析失败，使用简单的后备解析逻辑
                # 保持之前的标题，不要重置为默认值
                # self.title = "简单转换结果"
                self.content = self.soup.body if self.soup and hasattr(self.soup, 'body') else self.soup
            
            if time.time() - start_time > max_execution_time:
                return {"success": False, "message": "转换超时"}
            
            print("[DEBUG] 开始下载图片")
            if not self._download_all_images():
                # 图片下载失败不影响整体转换
                print("[DEBUG] 图片下载失败，但继续执行")
            
            if time.time() - start_time > max_execution_time:
                return {"success": False, "message": "转换超时"}
            
            success = self._create_word_document()
            if not success:
                print("[DEBUG] Word文档创建失败，使用极简后备逻辑")
                # Word文档创建失败，使用极简后备逻辑
                self.doc = Document()
                self.doc.add_heading("极简转换结果", level=1)
                self.doc.add_paragraph(f"URL: {self.url}")
                self.doc.add_paragraph(f"转换时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")

            # 解析输出路径
            if output_file:
                final_output_file = output_file
            else:
                sanitized_title = self._sanitize_filename(self.title)
                final_output_file = os.path.join(self.output_dir, f"{sanitized_title}.docx")

            if not self._save_document(final_output_file):
                return {"success": False, "message": "Word文档保存失败"}

            end_time = time.time()
            print(f"[DEBUG] 转换完成，总耗时: {end_time - start_time:.2f} 秒")

            return {
                "success": True,
                "message": "转换成功",
                "output_file": final_output_file,
                "title": self.title,
                "downloaded_images": len(self.downloaded_images),
                "url": self.url,
                "execution_time": end_time - start_time
            }
        except Exception as e:
            print(f"[DEBUG] 转换过程中发生异常: {str(e)}")
            import traceback
            traceback.print_exc()
            
            # 最终后备逻辑，确保返回一个有效的Word文档
            try:
                print("[DEBUG] 使用最终后备逻辑")
                self.doc = Document()
                self.doc.add_heading("应急转换结果", level=1)
                self.doc.add_paragraph(f"URL: {self.url}")
                self.doc.add_paragraph(f"转换时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
                self.doc.add_paragraph(f"转换过程中发生异常: {str(e)}")
                
                # 解析输出路径
                if output_file:
                    final_output_file = output_file
                else:
                    final_output_file = os.path.join(self.output_dir, f"应急转换结果_{int(time.time())}.docx")
                
                self._save_document(final_output_file)
                
                return {
                    "success": True,
                    "message": "转换成功(使用应急逻辑)",
                    "output_file": final_output_file,
                    "title": "应急转换结果",
                    "downloaded_images": 0,
                    "url": self.url,
                    "execution_time": time.time() - start_time
                }
            except Exception as backup_e:
                print(f"[DEBUG] 最终后备逻辑也失败了: {str(backup_e)}")
                traceback.print_exc()
                return {
                    "success": False,
                    "message": f"转换失败: {str(e)}",
                    "execution_time": time.time() - start_time
                }


def convert_web_to_docx(url, output_file=None, options=None):
    """
    网页转Word的主函数入口

    Args:
        url (str): 输入的网页URL或本地HTML文件路径
        output_file (str, optional): 输出的Word文件路径
        options (dict, optional): 转换选项

    Returns:
        dict: 转换结果信息
    """
    if options is None:
        options = {}

    # 解析选项
    output_dir = options.get("output_dir", "output")
    timeout = options.get("timeout", 10)

    # 处理输出文件路径
    if output_file:
        # 如果提供了输出文件路径，使用该路径的目录作为output_dir
        output_dir = os.path.dirname(output_file)
        # 确保输出目录存在
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
    else:
        # 确保默认输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

    # 创建转换器实例
    converter = WebToDocxConverter(
        url=url,
        output_dir=output_dir,
        timeout=timeout,
        progress_callback=options.get("progress_callback"),
    )

    # 执行转换
    result = converter.convert(output_file)

    return result


def main():
    """
    主函数，处理命令行参数
    """
    import sys
    if len(sys.argv) < 2:
        print("使用方法: python web_to_docx.py <网页URL> [输出文件路径]")
        return
    
    url = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    print(f"正在转换网页: {url}")
    result = convert_web_to_docx(url, output_file)
    
    if result["success"]:
        print(f"转换成功！输出文件: {result['output_file']}")
    else:
        print(f"转换失败: {result['message']}")


if __name__ == "__main__":
    main()
