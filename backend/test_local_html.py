#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
from docx import Document
from bs4 import BeautifulSoup

# 创建一个简单的HTML文件用于测试
def create_test_html():
    html_content = """<!DOCTYPE html>
<html>
<head>
    <title>测试页面</title>
</head>
<body>
    <h1>测试标题</h1>
    <p>这是一个测试段落。</p>
    <p>这是另一个测试段落。</p>
    <ul>
        <li>列表项1</li>
        <li>列表项2</li>
        <li>列表项3</li>
    </ul>
    <h2>二级标题</h2>
    <p>这是二级标题下的段落。</p>
</body>
</html>"""
    
    with open("test.html", "w", encoding="utf-8") as f:
        f.write(html_content)
    
    return "test.html"

# 直接测试HTML解析和Word文档生成
def test_local_html_to_docx():
    print("=== 直接测试HTML解析和Word文档生成 ===")
    
    # 创建测试HTML文件
    html_file = create_test_html()
    print(f"创建测试HTML文件: {html_file}")
    
    try:
        # 读取HTML内容
        with open(html_file, "r", encoding="utf-8") as f:
            html_content = f.read()
        
        print(f"HTML内容长度: {len(html_content)} 字符")
        
        # 解析HTML
        soup = BeautifulSoup(html_content, "html.parser")
        print("HTML解析完成")
        
        # 提取标题
        title = soup.title.string if soup.title else "测试文档"
        print(f"提取标题: {title}")
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, level=1)
        print("添加标题到文档")
        
        # 提取正文内容
        body = soup.body
        if body:
            # 遍历body的子元素
            for element in body.children:
                if element.name == "h1":
                    doc.add_heading(element.get_text(), level=1)
                elif element.name == "h2":
                    doc.add_heading(element.get_text(), level=2)
                elif element.name == "p":
                    doc.add_paragraph(element.get_text())
                elif element.name == "ul":
                    for li in element.find_all("li"):
                        doc.add_paragraph(li.get_text(), style="List Bullet")
                elif element.name == "ol":
                    for li in element.find_all("li"):
                        doc.add_paragraph(li.get_text(), style="List Number")
            
            print("添加正文内容到文档")
        
        # 保存文档
        output_file = "test_output.docx"
        doc.save(output_file)
        print(f"文档保存成功: {output_file}")
        print(f"文件大小: {os.path.getsize(output_file)} 字节")
        
        return True
        
    except Exception as e:
        print(f"测试失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理测试文件
        if os.path.exists(html_file):
            os.remove(html_file)

if __name__ == "__main__":
    test_local_html_to_docx()